import os
import re
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

from schemas import InfrastructureData, InstanceData, LoadBalancerData

# --- i18n: Embedded PT-BR and EN string tables ---

DOC_STRINGS = {
    "pt": {
        # --- Common Terms ---
        "doc.common.client": "Cliente",
        "doc.common.generation_date": "Data de Geração",
        "doc.common.toc": "Sumário",
        "doc.common.associated_hosts": "Hosts Associados: ",
        "doc.common.all": "Todas",
        "doc.common.managed_by_vg": "Gerenciado por VG",
        "doc.common.boot_volume": "Boot Volume",
        "doc.common.volume_group": "Volume Group",
        "doc.common.name": "Nome",
        "doc.common.state": "Estado",
        "doc.common.ad": "Availability Domain",
        "doc.common.member_count": "Total de Membros",
        "doc.common.backup_policy": "Política de Backup",
        "doc.common.cross_region_replication": "Replicação Cross-Region",
        "doc.common.enabled": "Habilitada",
        "doc.common.disabled": "Desabilitada",
        "doc.common.replication_target": "Destino da Replicação",
        "doc.common.public": "Público",
        "doc.common.private": "Privado",
        "doc.common.yes": "Sim",
        "doc.common.no": "Não",
        "doc.common.status": "Status",
        "doc.common.type": "Tipo",
        "doc.common.protocol": "Protocolo",
        "doc.common.port": "Porta",
        "doc.common.weight": "Peso",
        "doc.common.routing": "Roteamento",
        "doc.common.lifetime": "Lifetime (s)",
        "doc.common.disk_gb": "DISCO (GB)",
        "doc.common.memory_gb": "MEMÓRIA (GB)",
        # --- Headers ---
        "doc.headers.direction": "Direção",
        "doc.headers.protocol": "Protocolo",
        "doc.headers.ports": "Portas",
        "doc.headers.source_destination": "Origem/Destino",
        "doc.headers.description": "Descrição",
        "doc.headers.destination": "Destino",
        "doc.headers.target": "Alvo (Target)",
        "doc.headers.host": "HOST",
        "doc.headers.shape": "SHAPE",
        "doc.headers.ocpu": "OCPU",
        "doc.headers.memory_gb": "MEMÓRIA (GB)",
        "doc.headers.boot_vol_gb": "BOOT VOL (GB)",
        "doc.headers.os": "S.O.",
        "doc.headers.private_ip": "IP PRIVADO",
        "doc.headers.public_ip": "IP PÚBLICO",
        "doc.headers.source_host": "Host de Origem",
        "doc.headers.volume_name": "Nome do Volume",
        "doc.headers.size_gb": "Tamanho (GB)",
        "doc.headers.backup_policy": "Política de Backup",
        "doc.headers.associated_host": "Host Associado",
        "doc.headers.backup_policy_applied": "Política de Backup Aplicada",
        "doc.headers.group_members": "Membros do Grupo",
        "doc.headers.subnet_name": "Nome da Subnet",
        "doc.headers.cidr_block": "CIDR Block",
        "doc.headers.peering_status": "Status do Peering",
        "doc.headers.route_table": "Route Table",
        "doc.headers.advertised_cidr": "CIDR Anunciado",
        "doc.headers.cross_tenancy": "Cross-Tenancy",
        "doc.headers.cluster_name": "NOME DO CLUSTER",
        "doc.headers.k8s_version": "VERSÃO KUBERNETES",
        "doc.headers.vcn_associated": "VCN ASSOCIADA",
        "doc.headers.public_api_endpoint": "ENDPOINT API PÚBLICO",
        "doc.headers.private_api_endpoint": "ENDPOINT API PRIVADO",
        "doc.headers.nodepool_name": "NOME DO NODE POOL",
        "doc.headers.node_count": "QT. NODES",
        "doc.headers.subnet": "SUBNET",
        "doc.headers.lb_ip_addresses": "Endereços IP",
        "doc.headers.lb_listeners": "Listeners",
        "doc.headers.lb_backend_set_default": "Backend Set Padrão",
        "doc.headers.lb_backend_sets": "Backend Sets",
        "doc.headers.lb_backend_policy": "Política de Balanceamento",
        "doc.headers.lb_health_checker": "Health Checker",
        "doc.headers.lb_backend_name": "Nome do Backend",
        "doc.headers.lb_backend_ip": "IP",
        "doc.headers.lb_backend_port": "Porta",
        "doc.headers.lb_backend_weight": "Peso",
        "doc.headers.drg_attachment_name": "Nome do Anexo",
        "doc.headers.drg_attachment_type": "Tipo",
        "doc.headers.drg_route_table": "DRG Route Table",
        "doc.headers.rpc_name": "Nome da Conexão Remota",
        "doc.headers.rpc_peering_status": "Status do Peering",
        "doc.headers.cpe_name": "Nome do CPE",
        "doc.headers.cpe_ip": "Endereço IP",
        "doc.headers.cpe_vendor": "Fabricante",
        "doc.headers.vpn_connection_name": "Nome da Conexão",
        "doc.headers.vpn_cpe_associated": "CPE Associado",
        "doc.headers.vpn_drg_associated": "DRG Associado",
        "doc.headers.vpn_routing": "Roteamento",
        "doc.headers.vpn_tunnels": "Túneis",
        "doc.headers.vpn_oracle_ip": "IP Oracle",
        "doc.headers.vpn_cpe_ip": "IP do CPE",
        "doc.headers.vpn_ike_version": "Versão IKE",
        "doc.headers.vpn_bgp_oracle_asn": "ASN Oracle",
        "doc.headers.vpn_bgp_customer_asn": "ASN do Cliente",
        "doc.headers.vpn_bgp_oracle_ip": "IP do Túnel (Oracle)",
        "doc.headers.vpn_bgp_customer_ip": "IP do Túnel (Cliente)",
        "doc.headers.vpn_p1_auth": "Autenticação",
        "doc.headers.vpn_p1_encryption": "Criptografia",
        "doc.headers.vpn_p1_dh_group": "Grupo DH",
        "doc.headers.vpn_p1_lifetime": "Lifetime (s)",
        "doc.headers.vpn_p2_auth": "Autenticação",
        "doc.headers.vpn_p2_encryption": "Criptografia",
        "doc.headers.vpn_p2_lifetime": "Lifetime (s)",
        "doc.headers.responsible": "RESPONSÁVEL",
        "doc.headers.date": "DATA",
        # --- Headings ---
        "doc.headings.architecture": "Arquitetura e Escopo",
        "doc.headings.architecture_drawing": "Desenho da Arquitetura",
        "doc.headings.infra_config": "Configuração de Infraestrutura",
        "doc.headings.compute_instances": "Instâncias Computacionais",
        "doc.headings.volume_management": "Gerenciamento de Volumes e Backup",
        "doc.headings.attached_volumes": "Volumes Anexados (Block Volumes)",
        "doc.headings.backup_policies": "Políticas de Backup",
        "doc.headings.backup_policy_details": "Detalhes da Política de Backup CCM-7D",
        "doc.headings.standalone_volumes": "Block Volumes Desanexados (Sem Instância Associada)",
        "doc.descriptions.standalone_volumes": "Os volumes abaixo existem no compartimento mas não estão atualmente anexados a nenhuma instância computacional.",
        "doc.headings.volume_groups": "Volume Groups",
        "doc.headings.general_info": "Informações Gerais",
        "doc.headings.data_protection_validation": "Validação de Proteção de Dados",
        "doc.headings.vcn_topology": "Topologia de Rede Virtual (VCN)",
        "doc.headings.vcn_details": "Detalhes da VCN",
        "doc.headings.subnets": "Subnets",
        "doc.headings.security_lists": "Security Lists",
        "doc.headings.sl_rules": "Regras da Security List",
        "doc.headings.route_tables": "Route Tables",
        "doc.headings.rt_rules": "Regras da Route Table",
        "doc.headings.nsgs": "Network Security Groups (NSGs)",
        "doc.headings.nsg_rules": "Network Security Group",
        "doc.headings.lpgs": "Local Peering Gateways (LPGs)",
        "doc.headings.oke_config": "CONFIGURAÇÃO AMBIENTE OKE",
        "doc.headings.oke_cluster": "Cluster",
        "doc.headings.oke_nodepools": "Node Pools (Worker Nodes)",
        "doc.headings.load_balancers": "Load Balancers (LBaaS)",
        "doc.headings.lb_details": "Load Balancer",
        "doc.headings.lb_general_info": "Informações Gerais do Load Balancer",
        "doc.headings.lb_listeners": "Listeners",
        "doc.headings.lb_backend_sets": "Backend Sets",
        "doc.headings.lb_backend_set_config": "Configuração do Backend Set",
        "doc.headings.connectivity": "Conectividade Externa e Roteamento",
        "doc.headings.drgs": "Dynamic Routing Gateways (DRGs)",
        "doc.headings.drg_details": "DRG",
        "doc.headings.drg_attachments": "Anexos do DRG",
        "doc.headings.rpcs": "Conexões de Peering Remoto (RPCs)",
        "doc.headings.cpes": "Customer-Premises Equipment (CPEs)",
        "doc.headings.vpn_connections": "Conexões VPN IPSec",
        "doc.headings.vpn_summary": "Sumário de Conexões VPN",
        "doc.headings.vpn_details": "Detalhes das Conexões VPN",
        "doc.headings.vpn_connection_config": "Configuração da Conexão",
        "doc.headings.vpn_tunnel_info": "Informações do Túnel",
        "doc.headings.vpn_tunnel_details": "Túnel",
        "doc.headings.vpn_bgp_details": "Detalhes da Sessão BGP",
        "doc.headings.vpn_phase1": "Fase 1 (IKE)",
        "doc.headings.vpn_phase2": "Fase 2 (IPSec)",
        "doc.headings.vpn_config_note": "Nota de Configuração: ",
        "doc.headings.additional_configs": "Configurações Adicionais",
        "doc.headings.antivirus_config": "Configuração do Antivírus",
        "doc.headings.responsible": "RESPONSÁVEL",
        "doc.type.waf_report": "Documentação de WAF",
        "doc.identifier.waf_report": "WAF",
        "doc.headings.waf_config": "Configuração de Web Application Firewall (WAF)",
        "doc.headings.waf_policy_name": "Política WAF",
        "doc.headings.waf_integration": "Integração WAF",
        "doc.headings.waf_firewall": "Web App Firewall (Firewall)",
        "doc.headings.overview": "Visão Geral",
        "doc.headings.lb_integration": "Integração com Load Balancer",
        "doc.headings.network_infrastructure": "Infraestrutura de Rede",
        "doc.headings.waf_protection_config": "Configuração de Proteção WAF",
        "doc.headings.waf_actions": "Ações (Actions)",
        "doc.headings.waf_access_control": "Controle de Acesso",
        "doc.headings.waf_rate_limiting": "Rate Limiting",
        "doc.headings.waf_protection_rules": "Regras de Proteção",
        "doc.headers.protected_domains": "Domínios Protegidos",
        "doc.headers.capabilities": "Capabilities",
        "doc.headers.action": "Ação",
        "doc.headers.condition": "Condição",
        "doc.headers.certificates": "Certificados",
        # --- Firewall / LB binding table ---
        "doc.headers.firewall_name": "Nome do Firewall",
        "doc.headers.lb_name":       "Load Balancer",
        "doc.headers.lb_ip":         "Endereços IP",
        # --- Certificate detail fields ---
        "doc.headers.cert_org":           "Organização",
        "doc.headers.cert_key_algo":      "Algoritmo da Chave",
        "doc.headers.cert_sign_algo":     "Algoritmo de Assinatura",
        "doc.headers.cert_stages":        "Estágios",
        "doc.headers.cert_valid_from":    "Válido De",
        "doc.headers.cert_valid_to":      "Válido Até",
        "doc.headers.cert_deletion_date": "Data de Exclusão",
        "doc.headers.expiration": "Expiração",
        "doc.common.compartment": "Compartimento",
        "doc.common.region": "Região",
        "doc.common.time_created": "Data de Criação",
        "doc.common.code": "Código HTTP",
        "doc.messages.no_lb_association": "Esta política WAF não possui Load Balancer associado.",
        "doc.messages.no_network_found": "Nenhuma Infraestrutura de Rede encontrada.",
        "doc.messages.no_config_found": "Nenhuma configuração encontrada.",
        "doc.messages.no_certificates_found": "Nenhum certificado configurado.",
        "doc.headings.responsible_desc": "Responsável pelo preenchimento da documentação:",
        "doc.headings.instance_connectivity": "Conectividade de Rede da(s) Instância(s)",
        # --- Descriptions ---
        "doc.descriptions.instance_table": "A tabela a seguir detalha as configurações das instâncias computacionais no escopo.",
        "doc.descriptions.block_volume_table": "A tabela abaixo detalha os Block Volumes anexados às instâncias.",
        "doc.descriptions.backup_policy_table": "A tabela a seguir consolida as políticas de backup para todos os volumes (Boot e Block) das instâncias.",
        "doc.descriptions.backup_policy_details_1": "O Backup CCM-7D está configurado por meio de uma Backup Policy, garantindo a proteção contínua dos dados.",
        "doc.descriptions.backup_policy_details_2": "A política define um agendamento diário Incremental, executado à 01:00 (horário regional).",
        "doc.descriptions.backup_policy_details_3": "Os backups gerados possuem retenção de 7 dias.",
        "doc.descriptions.backup_policy_details_4": "Essa política garante que, em caso de falhas, seja possível recuperar o volume rapidamente.",
        "doc.descriptions.volume_groups": "Volume Groups são conjuntos de volumes (Boot e Block) que podem ser gerenciados como uma única unidade, especialmente para backups consistentes e replicação entre regiões.",
        "doc.descriptions.vpn_config_warning": "Os parâmetros de criptografia customizados para este túnel divergem das recomendações padrão da Oracle.",
        # --- Messages ---
        "doc.messages.no_rules_for": "Nenhuma regra configurada para este",
        "doc.messages.resource_not_provisioned": "Este recurso não está provisionado neste ambiente.",
        "doc.messages.no_vg_members": "Este Volume Group não possui membros.",
        "doc.messages.no_vcn_found": "Nenhuma infraestrutura de rede associada foi encontrada.",
        "doc.messages.no_subnet_found": "Nenhuma subnet encontrada nesta VCN.",
        "doc.messages.no_sl_found": "Nenhuma Security List encontrada nesta VCN.",
        "doc.messages.no_rt_found": "Nenhuma Route Table encontrada nesta VCN.",
        "doc.messages.no_nsg_found": "Nenhum Network Security Group encontrado para esta VCN.",
        "doc.messages.no_lpg_found": "Nenhum Local Peering Gateway encontrado nesta VCN.",
        "doc.messages.no_nodepool_found": "Nenhum Node Pool encontrado para este cluster.",
        "doc.messages.no_listener_found": "Nenhum Listener configurado.",
        "doc.messages.no_backend_set_found": "Nenhum Backend Set configurado.",
        "doc.messages.no_backend_found": "Nenhum backend configurado neste Backend Set.",
        "doc.messages.no_drg_found": "Nenhum Dynamic Routing Gateway encontrado neste compartimento.",
        "doc.messages.no_drg_attachment_found": "Nenhum anexo encontrado para este DRG.",
        "doc.messages.no_rpc_found": "Nenhuma Conexão de Peering Remoto (RPC) associada a este DRG.",
        "doc.messages.no_cpe_found": "Nenhum CPE encontrado neste compartimento.",
        "doc.messages.no_vpn_found": "Nenhuma Conexão VPN IPSec encontrada neste compartimento.",
        "doc.messages.no_tunnel_found": "Nenhum túnel encontrado para esta conexão.",
        "doc.messages.insert_image_error": "Erro ao inserir imagem",
        "doc.messages.no_rules_found": "Nenhuma regra configurada.",
        "doc.messages.no_routes_found": "Nenhuma rota configurada.",
        "doc.headers.vcn_name": "Nome da VCN",
        "doc.headers.vcn_cidr": "CIDR da VCN",
        "doc.headers.subnet_name": "Nome da Subnet",
        "doc.headers.subnet_cidr": "CIDR da Subnet",
        "doc.headers.destination_cidr": "CIDR de Destino",
        "doc.headers.common_name": "COMMON NAME",
        "doc.headers.sans": "SANs",
        "doc.headers.key_algo": "ALGORITMO DA CHAVE",
        "doc.headers.sig_algo": "ALGORITMO DE ASSINATURA",
        "doc.headers.valid_from": "VÁLIDO DE",
        "doc.headers.valid_to": "VÁLIDO ATÉ",
        "doc.headers.associated_listener": "ASSOCIADO AO LISTENER",
        "doc.common.none": "Nenhum",
        # --- Document Types ---
        "doc.type.full_infra": "Documentação de Infraestrutura",        "doc.type.new_host": "Documentação de Novo Host",
        "doc.type.kubernetes": "Documentação de Kubernetes (OKE)",
        "doc.type.default": "Documentação Técnica",
        "doc.identifier.full_infra": "Infraestrutura",
        "doc.identifier.new_host": "NovoHost",
        "doc.identifier.kubernetes": "Kubernetes",
        "doc.identifier.default": "Geral",
    },
    "en": {
        # --- Common Terms ---
        "doc.common.client": "Client",
        "doc.common.generation_date": "Generation Date",
        "doc.common.toc": "Table of Contents",
        "doc.common.associated_hosts": "Associated Hosts: ",
        "doc.common.all": "All",
        "doc.common.managed_by_vg": "Managed by VG",
        "doc.common.boot_volume": "Boot Volume",
        "doc.common.volume_group": "Volume Group",
        "doc.common.name": "Name",
        "doc.common.state": "State",
        "doc.common.ad": "Availability Domain",
        "doc.common.member_count": "Total Members",
        "doc.common.backup_policy": "Backup Policy",
        "doc.common.cross_region_replication": "Cross-Region Replication",
        "doc.common.enabled": "Enabled",
        "doc.common.disabled": "Disabled",
        "doc.common.replication_target": "Replication Target",
        "doc.common.public": "Public",
        "doc.common.private": "Private",
        "doc.common.yes": "Yes",
        "doc.common.no": "No",
        "doc.common.status": "Status",
        "doc.common.type": "Type",
        "doc.common.protocol": "Protocol",
        "doc.common.port": "Port",
        "doc.common.weight": "Weight",
        "doc.common.routing": "Routing",
        "doc.common.lifetime": "Lifetime (s)",
        "doc.common.disk_gb": "DISK (GB)",
        "doc.common.memory_gb": "MEMORY (GB)",
        # --- Headers ---
        "doc.headers.direction": "Direction",
        "doc.headers.protocol": "Protocol",
        "doc.headers.ports": "Ports",
        "doc.headers.source_destination": "Source/Destination",
        "doc.headers.description": "Description",
        "doc.headers.destination": "Destination",
        "doc.headers.target": "Target",
        "doc.headers.host": "HOST",
        "doc.headers.shape": "SHAPE",
        "doc.headers.ocpu": "OCPU",
        "doc.headers.memory_gb": "MEMORY (GB)",
        "doc.headers.boot_vol_gb": "BOOT VOL (GB)",
        "doc.headers.os": "OS",
        "doc.headers.private_ip": "PRIVATE IP",
        "doc.headers.public_ip": "PUBLIC IP",
        "doc.headers.source_host": "Source Host",
        "doc.headers.volume_name": "Volume Name",
        "doc.headers.size_gb": "Size (GB)",
        "doc.headers.backup_policy": "Backup Policy",
        "doc.headers.associated_host": "Associated Host",
        "doc.headers.backup_policy_applied": "Backup Policy Applied",
        "doc.headers.group_members": "Group Members",
        "doc.headers.subnet_name": "Subnet Name",
        "doc.headers.cidr_block": "CIDR Block",
        "doc.headers.peering_status": "Peering Status",
        "doc.headers.route_table": "Route Table",
        "doc.headers.advertised_cidr": "Advertised CIDR",
        "doc.headers.cross_tenancy": "Cross-Tenancy",
        "doc.headers.cluster_name": "CLUSTER NAME",
        "doc.headers.k8s_version": "KUBERNETES VERSION",
        "doc.headers.vcn_associated": "ASSOCIATED VCN",
        "doc.headers.public_api_endpoint": "PUBLIC API ENDPOINT",
        "doc.headers.private_api_endpoint": "PRIVATE API ENDPOINT",
        "doc.headers.nodepool_name": "NODE POOL NAME",
        "doc.headers.node_count": "NODE QTY",
        "doc.headers.subnet": "SUBNET",
        "doc.headers.lb_ip_addresses": "IP Addresses",
        "doc.headers.lb_listeners": "Listeners",
        "doc.headers.lb_backend_set_default": "Default Backend Set",
        "doc.headers.lb_backend_sets": "Backend Sets",
        "doc.headers.lb_backend_policy": "Balancing Policy",
        "doc.headers.lb_health_checker": "Health Checker",
        "doc.headers.lb_backend_name": "Backend Name",
        "doc.headers.lb_backend_ip": "IP",
        "doc.headers.lb_backend_port": "Port",
        "doc.headers.lb_backend_weight": "Weight",
        "doc.headers.drg_attachment_name": "Attachment Name",
        "doc.headers.drg_attachment_type": "Type",
        "doc.headers.drg_route_table": "DRG Route Table",
        "doc.headers.rpc_name": "Remote Connection Name",
        "doc.headers.rpc_peering_status": "Peering Status",
        "doc.headers.cpe_name": "CPE Name",
        "doc.headers.cpe_ip": "IP Address",
        "doc.headers.cpe_vendor": "Vendor",
        "doc.headers.vpn_connection_name": "Connection Name",
        "doc.headers.vpn_cpe_associated": "Associated CPE",
        "doc.headers.vpn_drg_associated": "Associated DRG",
        "doc.headers.vpn_routing": "Routing",
        "doc.headers.vpn_tunnels": "Tunnels",
        "doc.headers.vpn_oracle_ip": "Oracle IP",
        "doc.headers.vpn_cpe_ip": "CPE IP",
        "doc.headers.vpn_ike_version": "IKE Version",
        "doc.headers.vpn_bgp_oracle_asn": "Oracle ASN",
        "doc.headers.vpn_bgp_customer_asn": "Customer ASN",
        "doc.headers.vpn_bgp_oracle_ip": "Tunnel IP (Oracle)",
        "doc.headers.vpn_bgp_customer_ip": "Tunnel IP (Customer)",
        "doc.headers.vpn_p1_auth": "Authentication",
        "doc.headers.vpn_p1_encryption": "Encryption",
        "doc.headers.vpn_p1_dh_group": "DH Group",
        "doc.headers.vpn_p1_lifetime": "Lifetime (s)",
        "doc.headers.vpn_p2_auth": "Authentication",
        "doc.headers.vpn_p2_encryption": "Encryption",
        "doc.headers.vpn_p2_lifetime": "Lifetime (s)",
        "doc.headers.responsible": "RESPONSIBLE",
        "doc.headers.date": "DATE",
        # --- Headings ---
        "doc.headings.architecture": "Architecture and Scope",
        "doc.headings.architecture_drawing": "Architecture Drawing",
        "doc.headings.infra_config": "Infrastructure Configuration",
        "doc.headings.compute_instances": "Compute Instances",
        "doc.headings.volume_management": "Volume and Backup Management",
        "doc.headings.attached_volumes": "Attached Block Volumes",
        "doc.headings.backup_policies": "Backup Policies",
        "doc.headings.backup_policy_details": "CCM-7D Backup Policy Details",
        "doc.headings.standalone_volumes": "Unattached Block Volumes (No Associated Instance)",
        "doc.descriptions.standalone_volumes": "The volumes below exist in the compartment but are not currently attached to any compute instance.",
        "doc.headings.volume_groups": "Volume Groups",
        "doc.headings.general_info": "General Information",
        "doc.headings.data_protection_validation": "Data Protection Validation",
        "doc.headings.vcn_topology": "Virtual Network Topology (VCN)",
        "doc.headings.vcn_details": "VCN Details",
        "doc.headings.subnets": "Subnets",
        "doc.headings.security_lists": "Security Lists",
        "doc.headings.sl_rules": "Security List Rules",
        "doc.headings.route_tables": "Route Tables",
        "doc.headings.rt_rules": "Route Table Rules",
        "doc.headings.nsgs": "Network Security Groups (NSGs)",
        "doc.headings.nsg_rules": "Network Security Group",
        "doc.headings.lpgs": "Local Peering Gateways (LPGs)",
        "doc.headings.oke_config": "OKE ENVIRONMENT CONFIGURATION",
        "doc.headings.oke_cluster": "Cluster",
        "doc.headings.oke_nodepools": "Node Pools (Worker Nodes)",
        "doc.headings.load_balancers": "Load Balancers (LBaaS)",
        "doc.headings.lb_details": "Load Balancer",
        "doc.headings.lb_general_info": "General Load Balancer Information",
        "doc.headings.lb_listeners": "Listeners",
        "doc.headings.lb_backend_sets": "Backend Sets",
        "doc.headings.lb_backend_set_config": "Backend Set Configuration",
        "doc.headings.connectivity": "External Connectivity and Routing",
        "doc.headings.drgs": "Dynamic Routing Gateways (DRGs)",
        "doc.headings.drg_details": "DRG",
        "doc.headings.drg_attachments": "DRG Attachments",
        "doc.headings.rpcs": "Remote Peering Connections (RPCs)",
        "doc.headings.cpes": "Customer-Premises Equipment (CPEs)",
        "doc.headings.vpn_connections": "IPSec VPN Connections",
        "doc.headings.vpn_summary": "VPN Connections Summary",
        "doc.headings.vpn_details": "VPN Connection Details",
        "doc.headings.vpn_connection_config": "Connection Configuration",
        "doc.headings.vpn_tunnel_info": "Tunnel Information",
        "doc.headings.vpn_tunnel_details": "Tunnel",
        "doc.headings.vpn_bgp_details": "BGP Session Details",
        "doc.headings.vpn_phase1": "Phase 1 (IKE)",
        "doc.headings.vpn_phase2": "Phase 2 (IPSec)",
        "doc.headings.vpn_config_note": "Configuration Note: ",
        "doc.headings.additional_configs": "Additional Configurations",
        "doc.headings.antivirus_config": "Antivirus Configuration",
        "doc.headings.responsible": "RESPONSIBLE",
        "doc.type.waf_report": "WAF Documentation",
        "doc.identifier.waf_report": "WAF",
        "doc.headings.waf_config": "Web Application Firewall (WAF) Configuration",
        "doc.headings.waf_policy_name": "WAF Policy",
        "doc.headings.waf_integration": "WAF Integration",
        "doc.headings.waf_firewall": "Web App Firewall (Firewall)",
        "doc.headings.overview": "Overview",
        "doc.headings.lb_integration": "Load Balancer Integration",
        "doc.headings.network_infrastructure": "Network Infrastructure",
        "doc.headings.waf_protection_config": "WAF Protection Configuration",
        "doc.headings.waf_actions": "Actions",
        "doc.headings.waf_access_control": "Access Control",
        "doc.headings.waf_rate_limiting": "Rate Limiting",
        "doc.headings.waf_protection_rules": "Protection Rules",
        "doc.headers.protected_domains": "Protected Domains",
        "doc.headers.capabilities": "Capabilities",
        "doc.headers.action": "Action",
        "doc.headers.condition": "Condition",
        "doc.headers.certificates": "Certificates",
        # --- Firewall / LB binding table ---
        "doc.headers.firewall_name": "Firewall Name",
        "doc.headers.lb_name":       "Load Balancer",
        "doc.headers.lb_ip":         "IP Addresses",
        # --- Certificate detail fields ---
        "doc.headers.cert_org":           "Organization",
        "doc.headers.cert_key_algo":      "Key Algorithm",
        "doc.headers.cert_sign_algo":     "Signature Algorithm",
        "doc.headers.cert_stages":        "Stages",
        "doc.headers.cert_valid_from":    "Valid From",
        "doc.headers.cert_valid_to":      "Valid To",
        "doc.headers.cert_deletion_date": "Deletion Date",
        "doc.headers.expiration": "Expiration",
        "doc.common.compartment": "Compartment",
        "doc.common.region": "Region",
        "doc.common.time_created": "Time Created",
        "doc.common.code": "HTTP Code",
        "doc.messages.no_lb_association": "This WAF policy has no associated Load Balancer.",
        "doc.messages.no_network_found": "No Network Infrastructure found.",
        "doc.messages.no_config_found": "No configuration found.",
        "doc.messages.no_certificates_found": "No certificates configured.",
        "doc.headings.responsible_desc": "Responsible for filling out the documentation:",
        "doc.headings.instance_connectivity": "Instance Network Connectivity",
        # --- Descriptions ---
        "doc.descriptions.instance_table": "The following table details the configurations of the compute instances in scope.",
        "doc.descriptions.block_volume_table": "The table below details the Block Volumes attached to the instances.",
        "doc.descriptions.backup_policy_table": "The following table consolidates the backup policies for all volumes (Boot and Block) of the instances.",
        "doc.descriptions.backup_policy_details_1": "The CCM-7D Backup is configured via a Backup Policy, ensuring continuous data protection.",
        "doc.descriptions.backup_policy_details_2": "The policy defines a daily Incremental schedule, executed at 01:00 (regional time).",
        "doc.descriptions.backup_policy_details_3": "The generated backups have a retention of 7 days.",
        "doc.descriptions.backup_policy_details_4": "This policy ensures that in case of failures, the volume can be recovered quickly.",
        "doc.descriptions.volume_groups": "Volume Groups are sets of volumes (Boot and Block) that can be managed as a single unit, especially for consistent backups and cross-region replication.",
        "doc.descriptions.vpn_config_warning": "The custom encryption parameters for this tunnel differ from Oracle's standard recommendations.",
        # --- Messages ---
        "doc.messages.no_rules_for": "No rules configured for this",
        "doc.messages.resource_not_provisioned": "This resource is not provisioned in this environment.",
        "doc.messages.no_vg_members": "This Volume Group has no members.",
        "doc.messages.no_vcn_found": "No associated network infrastructure found.",
        "doc.messages.no_subnet_found": "No subnet found in this VCN.",
        "doc.messages.no_sl_found": "No Security List found in this VCN.",
        "doc.messages.no_rt_found": "No Route Table found in this VCN.",
        "doc.messages.no_nsg_found": "No Network Security Group found for this VCN.",
        "doc.messages.no_lpg_found": "No Local Peering Gateway found in this VCN.",
        "doc.messages.no_nodepool_found": "No Node Pool found for this cluster.",
        "doc.messages.no_listener_found": "No Listener configured.",
        "doc.messages.no_backend_set_found": "No Backend Set configured.",
        "doc.messages.no_backend_found": "No backend configured in this Backend Set.",
        "doc.messages.no_drg_found": "No Dynamic Routing Gateway found in this compartment.",
        "doc.messages.no_drg_attachment_found": "No attachments found for this DRG.",
        "doc.messages.no_rpc_found": "No Remote Peering Connection (RPC) associated with this DRG.",
        "doc.messages.no_cpe_found": "No CPE found in this compartment.",
        "doc.messages.no_vpn_found": "No IPSec VPN Connection found in this compartment.",
        "doc.messages.no_tunnel_found": "No tunnels found for this connection.",
        "doc.messages.insert_image_error": "Error inserting image",
        "doc.messages.no_rules_found": "No rules configured.",
        "doc.messages.no_routes_found": "No routes configured.",
        "doc.headers.vcn_name": "VCN Name",
        "doc.headers.vcn_cidr": "VCN CIDR",
        "doc.headers.subnet_name": "Subnet Name",
        "doc.headers.subnet_cidr": "Subnet CIDR",
        "doc.headers.destination_cidr": "Destination CIDR",
        "doc.headers.common_name": "COMMON NAME",
        "doc.headers.sans": "SANs",
        "doc.headers.key_algo": "KEY ALGORITHM",
        "doc.headers.sig_algo": "SIGNATURE ALGORITHM",
        "doc.headers.valid_from": "VALID FROM",
        "doc.headers.valid_to": "VALID TO",
        "doc.headers.associated_listener": "ASSOCIATED LISTENER",
        "doc.common.none": "None",
        # --- Document Types ---
        "doc.type.full_infra": "Infrastructure Documentation",
        "doc.type.new_host": "New Host Documentation",
        "doc.type.kubernetes": "Kubernetes (OKE) Documentation",
        "doc.type.default": "Technical Documentation",
        "doc.identifier.full_infra": "Infrastructure",
        "doc.identifier.new_host": "NewHost",
        "doc.identifier.kubernetes": "Kubernetes",
        "doc.identifier.default": "General",
    },
}


def t(key: str, lang: str) -> str:
    """Fetches a translation string based on the key and language."""
    # Fallback chain: requested lang → "pt" → the key itself.
    lang_to_use = lang if lang in DOC_STRINGS else "pt"
    strings = DOC_STRINGS.get(lang_to_use, {})
    return strings.get(key, DOC_STRINGS.get("pt", {}).get(key, key))


# --- Table of Contents and Hyperlink Helpers ---
def _define_toc_styles(document: Document):
    """Creates and configures 'TOC 1-3' styles with correct indentation if they don't exist."""
    styles = document.styles
    if "TOC 1" not in styles:
        style = styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        font = style.font
        font.name = "Gotham Light"
        font.size = Pt(11)
        p_format = style.paragraph_format
        p_format.left_indent = Inches(0)
        p_format.space_after = Pt(4)
    if "TOC 2" not in styles:
        style = styles.add_style("TOC 2", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        font = style.font
        font.name = "Gotham Light"
        font.size = Pt(11)
        p_format = style.paragraph_format
        p_format.left_indent = Inches(0.25)
        p_format.space_after = Pt(4)
    if "TOC 3" not in styles:
        style = styles.add_style("TOC 3", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        font = style.font
        font.name = "Gotham Light"
        font.size = Pt(11)
        p_format = style.paragraph_format
        p_format.left_indent = Inches(0.5)
        p_format.space_after = Pt(4)


def add_hyperlink(paragraph: Paragraph, text: str, url: str):
    """Adds an external hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    rPr.append(r_style)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_bookmark(paragraph: Paragraph, bookmark_name: str):
    """Adds a bookmark anchor to a paragraph, which can be linked to internally."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    bookmark_id = str(abs(hash(bookmark_name)) % (10**8))
    start_tag = OxmlElement("w:bookmarkStart")
    start_tag.set(qn("w:id"), bookmark_id)
    start_tag.set(qn("w:name"), bookmark_name)
    run._r.addprevious(start_tag)
    end_tag = OxmlElement("w:bookmarkEnd")
    end_tag.set(qn("w:id"), bookmark_id)
    run._r.addnext(end_tag)


def _add_internal_hyperlink(paragraph: Paragraph, text: str, anchor_name: str):
    """Adds an internal hyperlink that jumps to a bookmark within the document."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor_name)
    sub_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    sub_run.append(rPr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    sub_run.append(text_element)
    hyperlink.append(sub_run)
    paragraph._p.append(hyperlink)


def _add_and_bookmark_heading(
    document: Document,
    text: str,
    level: int,
    toc_list: list,
    counters: Dict[int, int],
):
    """Adds a numbered heading, creates a bookmark for it, and adds it to the TOC list."""
    if level not in counters:
        counters[level] = 0
    counters[level] += 1
    for deeper_level in range(level + 1, 6):
        if deeper_level in counters:
            counters[deeper_level] = 0

    number_parts = [str(counters.get(i, 0)) for i in range(1, level + 1)]
    number_str = ".".join(number_parts)
    final_text = f"{number_str} {text}"
    heading = document.add_paragraph(final_text, style=f"Heading {level}")

    clean_text = re.sub(r"[^A-Za-z0-9_]", "", final_text.replace(" ", "_"))
    bookmark_name = f"_Toc_{clean_text}_{len(toc_list)}"
    _add_bookmark(heading, bookmark_name)
    toc_list.append((final_text, bookmark_name, level))
    return heading


# --- Table Styling Helpers ---
def _shade_cell(cell, color="4472C4"):
    """Applies a background color shading to a table cell."""
    shading_xml = r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color)
    shading_elm = parse_xml(shading_xml)
    cell._tc.get_or_add_tcPr().append(shading_elm)


# Row background color by lifecycle state (hex without '#').
# TERMINATED       → light red    (resource deleted)
# STOPPED          → light orange (resource stopped)
# PENDING_DELETION → light yellow (scheduled for removal)
# All other states → no highlight
_STATE_ROW_COLORS: Dict[str, Optional[str]] = {
    "TERMINATED":      "FADBD8",   # light red
    "TERMINATING":     "FADBD8",   # light red
    "STOPPED":         "FDEBD0",   # light orange
    "STOPPING":        "FDEBD0",   # light orange
    "PENDING_DELETION":"FEF9C3",   # light yellow
}

_STATE_TEXT_COLORS: Dict[str, str] = {
    "TERMINATED":      "922B21",   # dark red text
    "TERMINATING":     "922B21",
    "STOPPED":         "784212",   # dark brown/orange text
    "STOPPING":        "784212",
    "PENDING_DELETION":"7D6608",   # dark yellow/olive text
}


def _style_row_by_state(row, state_raw: str) -> None:
    """
    Applies background color and text color to all cells in a row based on the
    resource lifecycle state. TERMINATED rows also get strikethrough to visually
    reinforce that the resource has been deleted.
    """
    state_key = state_raw.upper()
    bg_color   = _STATE_ROW_COLORS.get(state_key)
    text_color = _STATE_TEXT_COLORS.get(state_key)
    is_terminated = state_key in ("TERMINATED", "TERMINATING")

    if not bg_color and not text_color:
        return  # Active resources — no styling needed

    for cell in row.cells:
        if bg_color:
            _shade_cell(cell, bg_color)

        for para in cell.paragraphs:
            for run in para.runs:
                if text_color:
                    run.font.color.rgb = RGBColor.from_string(text_color)
                if is_terminated:
                    run.font.strike = True
            # If the paragraph has text but no runs (text set directly on the cell),
            # clear and re-add it with the correct formatting applied.
            if para.text and not para.runs:
                text = para.text
                para.clear()
                run = para.add_run(text)
                if text_color:
                    run.font.color.rgb = RGBColor.from_string(text_color)
                if is_terminated:
                    run.font.strike = True




def _style_cert_kv_table_by_state(document: Document, state: str) -> None:
    """
    Locates the last table added to the document (certificate key-value table)
    and applies background color to all data rows based on the certificate
    lifecycle state. The title row keeps the default blue color.
    """
    state_key  = state.upper()
    bg_color   = _STATE_ROW_COLORS.get(state_key)
    text_color = _STATE_TEXT_COLORS.get(state_key)
    if not bg_color and not text_color:
        return
    tables = document.tables
    if not tables:
        return
    table = tables[-1]
    for row in table.rows[1:]:
        for cell in row.cells:
            if bg_color:
                _shade_cell(cell, bg_color)
            for para in cell.paragraphs:
                for run in para.runs:
                    if text_color:
                        run.font.color.rgb = RGBColor.from_string(text_color)
                if para.text and not para.runs:
                    text = para.text
                    para.clear()
                    run = para.add_run(text)
                    if text_color:
                        run.font.color.rgb = RGBColor.from_string(text_color)

def _style_table_headers(table, headers, shade_color="4472C4"):
    """Applies standard formatting to a table's header row."""
    for i, header_text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text.upper())
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        _shade_cell(cell, shade_color)


def _create_titled_key_value_table(
    document, title, data_dict, col_widths=(Inches(2.0), Inches(4.5))
):
    """Creates a two-column table with a merged, styled title header."""
    table = document.add_table(rows=len(data_dict) + 1, cols=2, style="Table Grid")
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = col_widths[0]
    table.columns[1].width = col_widths[1]

    title_cell = table.cell(0, 0).merge(table.cell(0, 1))
    title_cell.text = ""
    p = title_cell.paragraphs[0]
    run = p.add_run(title.upper())
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("FFFFFF")
    _shade_cell(title_cell)

    for i, (key, value) in enumerate(data_dict.items(), 1):
        key_cell = table.cell(i, 0)
        key_cell.text = key
        key_cell.paragraphs[0].runs[0].font.bold = True
        table.cell(i, 1).text = str(value) if value else "N/A"
    document.add_paragraph()


# --- Content Formatting Helpers ---
def _add_network_resource_details(
    document: Document,
    resource_title: str,
    resource_name: str,
    rules: List[Any],
    associated_hosts: Optional[List[str]],
    rule_type: str,
    lang: str,
):
    """Adds a standardized section for a network resource (SL, NSG, RT) and its rules."""
    p = document.add_paragraph()
    p.add_run(f"{resource_title}: {resource_name}").bold = True

    if associated_hosts:
        p_hosts = document.add_paragraph()
        p_hosts.add_run(t("doc.common.associated_hosts", lang)).bold = True
        p_hosts.add_run(", ".join(sorted(associated_hosts)))

    if rule_type == "security":
        headers = [
            t("doc.headers.direction", lang),
            t("doc.headers.protocol", lang),
            t("doc.headers.ports", lang),
            t("doc.headers.source_destination", lang),
            t("doc.headers.description", lang),
        ]
    else:  # route
        headers = [
            t("doc.headers.destination", lang),
            t("doc.headers.target", lang),
            t("doc.headers.description", lang),
        ]

    if rules:
        table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
        _style_table_headers(table, headers)
        for rule in rules:
            cells = table.add_row().cells
            if rule_type == "security":
                cells[0].text = rule.direction
                cells[1].text = rule.protocol
                cells[2].text = rule.ports or t("doc.common.all", lang)
                cells[3].text = rule.source_or_destination or "N/A"
                cells[4].text = rule.description or ""
            else:  # route
                target_str = rule.target
                # If the target is a raw OCID (not resolved by the backend), apply generic label.
                if target_str and target_str.startswith("ocid1."):
                    target_lower = target_str.lower()
                    if "internetgateway" in target_lower:
                        target_str = "Internet Gateway"
                    elif "natgateway" in target_lower:
                        target_str = "NAT Gateway"
                    elif "servicegateway" in target_lower:
                        target_str = "Service Gateway"
                    elif "drg" in target_lower:
                        target_str = "Dynamic Routing Gateway (DRG)"
                    elif "localpeeringgateway" in target_lower:
                        target_str = "Local Peering Gateway (LPG)"
                    elif "privateip" in target_lower:
                        target_str = "Private IP"

                cells[0].text = rule.destination
                cells[1].text = target_str
                cells[2].text = rule.description or ""
    else:
        document.add_paragraph(f"{t('doc.messages.no_rules_for', lang)} {resource_title}.")

    document.add_paragraph()


# --- Document Section Generators ---
def _add_instances_table(document: Document, instances: List[InstanceData], lang: str):
    """Adds the main summary table of compute instances to the document."""
    if not instances:
        return
    # State label map for instances — normalises TERMINATED to a display-friendly string.
    state_label_map = {
        "RUNNING":    "Ativo"         if lang == "pt" else "Running",
        "STOPPED":    "Parado"        if lang == "pt" else "Stopped",
        "TERMINATED": "Excluído"      if lang == "pt" else "Terminated",
        "TERMINATING":"Excluindo..."  if lang == "pt" else "Terminating...",
    }
    document.add_paragraph(t("doc.descriptions.instance_table", lang))
    headers = [
        t("doc.headers.host", lang),
        t("doc.common.state", lang),
        t("doc.headers.shape", lang),
        t("doc.headers.ocpu", lang),
        t("doc.headers.memory_gb", lang),
        t("doc.headers.boot_vol_gb", lang),
        t("doc.headers.os", lang),
        t("doc.headers.private_ip", lang),
        t("doc.headers.public_ip", lang),
    ]
    table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
    _style_table_headers(table, headers)
    for data in instances:
        state_raw = (data.lifecycle_state or "").upper()
        state_display = state_label_map.get(state_raw, data.lifecycle_state or "N/A")
        cells = table.add_row().cells
        cells[0].text = data.host_name
        cells[1].text = state_display
        cells[2].text = data.shape
        cells[3].text = str(data.ocpus)
        cells[4].text = str(data.memory)
        cells[5].text = str(data.boot_volume_gb)
        cells[6].text = data.os_name
        cells[7].text = data.private_ip
        cells[8].text = data.public_ip or "N/A"
        # Apply alert color to the row based on instance lifecycle state.
        _style_row_by_state(table.rows[-1], state_raw)
    document.add_paragraph()


def _add_volume_and_backup_section(
    document: Document,
    infra_data: InfrastructureData,
    toc_list: list,
    counters: Dict[int, int],
    lang: str,
):
    """Adds the section covering attached block volumes and backup policies."""
    instances = infra_data.instances
    if not instances:
        return
    volume_to_vg_map = {
        vol_id: vg.display_name
        for vg in getattr(infra_data, "volume_groups", [])
        if hasattr(vg, "member_ids")
        for vol_id in vg.member_ids
    }
    _add_and_bookmark_heading(
        document, t("doc.headings.volume_management", lang), 2, toc_list, counters
    )
    _add_and_bookmark_heading(
        document, t("doc.headings.attached_volumes", lang), 3, toc_list, counters
    )
    all_block_volumes = [vol for data in instances for vol in data.block_volumes]
    if all_block_volumes:
        document.add_paragraph(t("doc.descriptions.block_volume_table", lang))
        headers = [
            t("doc.headers.source_host", lang),
            t("doc.headers.volume_name", lang),
            t("doc.headers.size_gb", lang),
            t("doc.headers.backup_policy", lang),
        ]
        table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
        _style_table_headers(table, headers)
        for data in instances:
            for vol in data.block_volumes:
                cells = table.add_row().cells
                cells[0].text = data.host_name
                cells[1].text = vol.display_name
                cells[2].text = str(int(vol.size_in_gbs))
                if vol.id in volume_to_vg_map:
                    cells[3].text = f"{t('doc.common.managed_by_vg', lang)}: {volume_to_vg_map.get(vol.id)}"
                else:
                    cells[3].text = vol.backup_policy_name
        document.add_paragraph()
    else:
        document.add_paragraph(t("doc.messages.resource_not_provisioned", lang))
    _add_and_bookmark_heading(
        document, t("doc.headings.backup_policies", lang), 3, toc_list, counters
    )
    document.add_paragraph(t("doc.descriptions.backup_policy_table", lang))

    # State label map for instances in the backup table.
    state_label_map = {
        "RUNNING":    "Ativo"    if lang == "pt" else "Running",
        "STOPPED":    "Parado"   if lang == "pt" else "Stopped",
        "TERMINATED": "Excluído" if lang == "pt" else "Terminated",
    }

    headers = [
        t("doc.headers.associated_host", lang),
        t("doc.common.state",            lang),
        t("doc.headers.volume_name",     lang),
        t("doc.headers.backup_policy_applied", lang),
    ]
    table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
    _style_table_headers(table, headers)
    for instance in instances:
        state_raw     = (instance.lifecycle_state or "").upper()
        state_display = state_label_map.get(state_raw, instance.lifecycle_state or "N/A")

        boot_cells = table.add_row().cells
        boot_cells[0].text = instance.host_name
        boot_cells[1].text = state_display
        boot_cells[2].text = t("doc.common.boot_volume", lang)
        if instance.boot_volume_id in volume_to_vg_map:
            boot_cells[3].text = (
                f"{t('doc.common.managed_by_vg', lang)}: {volume_to_vg_map.get(instance.boot_volume_id)}"
            )
        else:
            boot_cells[3].text = instance.backup_policy_name
        # Highlight row if the instance is terminated or stopped.
        _style_row_by_state(table.rows[-1], state_raw)

        for vol in instance.block_volumes:
            block_cells = table.add_row().cells
            block_cells[0].text = instance.host_name
            block_cells[1].text = state_display
            block_cells[2].text = vol.display_name
            if vol.id in volume_to_vg_map:
                block_cells[3].text = f"{t('doc.common.managed_by_vg', lang)}: {volume_to_vg_map.get(vol.id)}"
            else:
                block_cells[3].text = vol.backup_policy_name
            _style_row_by_state(table.rows[-1], state_raw)
    document.add_paragraph()

    has_instance_policy = any(
        i.backup_policy_name == "CCM-7D"
        or any(bv.backup_policy_name == "CCM-7D" for bv in i.block_volumes)
        for i in instances
    )
    has_vg_policy = False
    if hasattr(infra_data, "volume_groups") and infra_data.volume_groups:
        has_vg_policy = any(
            vg.validation.policy_name == "CCM-7D" for vg in infra_data.volume_groups
        )

    if has_instance_policy or has_vg_policy:
        _add_and_bookmark_heading(
            document,
            t("doc.headings.backup_policy_details", lang),
            4,
            toc_list,
            counters,
        )
        document.add_paragraph(t("doc.descriptions.backup_policy_details_1", lang))
        document.add_paragraph(t("doc.descriptions.backup_policy_details_2", lang))
        document.add_paragraph(t("doc.descriptions.backup_policy_details_3", lang))
        document.add_paragraph(t("doc.descriptions.backup_policy_details_4", lang))
    document.add_paragraph()

    # Standalone (unattached) block volumes
    standalone_vols = getattr(infra_data, "standalone_volumes", [])
    if standalone_vols:
        _add_and_bookmark_heading(
            document, t("doc.headings.standalone_volumes", lang), 3, toc_list, counters
        )
        document.add_paragraph(t("doc.descriptions.standalone_volumes", lang))
        headers_sv = [
            t("doc.headers.volume_name", lang),
            t("doc.headers.size_gb", lang),
            t("doc.common.state", lang),
            t("doc.headers.backup_policy", lang),
            "Availability Domain",
        ]
        table_sv = document.add_table(rows=1, cols=len(headers_sv), style="Table Grid")
        _style_table_headers(table_sv, headers_sv)
        for sv in sorted(standalone_vols, key=lambda v: v.display_name):
            sv_state_raw = (sv.lifecycle_state or "AVAILABLE").upper()
            sv_state_map = {
                "ACTIVE": t("doc.common.state_active", lang) if lang == "pt" else "Active",
                "AVAILABLE": t("doc.common.state_active", lang) if lang == "pt" else "Active",
                "TERMINATED": "Excluído" if lang == "pt" else "Deleted",
                "TERMINATING": "Excluindo..." if lang == "pt" else "Deleting...",
                "STOPPED": "Parado" if lang == "pt" else "Stopped",
                "STOPPING": "Parando" if lang == "pt" else "Stopping",
            }
            sv_state_display = sv_state_map.get(sv_state_raw, sv.lifecycle_state or "N/A")
            cells = table_sv.add_row().cells
            cells[0].text = sv.display_name
            cells[1].text = str(int(sv.size_in_gbs))
            cells[2].text = sv_state_display
            cells[3].text = sv.backup_policy_name
            cells[4].text = sv.availability_domain or "N/A"
            _style_row_by_state(table_sv.rows[-1], sv_state_raw)
        document.add_paragraph()


def _add_volume_groups_section(
    document: Document,
    infra_data: InfrastructureData,
    toc_list: list,
    counters: Dict[int, int],
    lang: str,
):
    """Adds the Volume Groups details section to the document."""
    if not infra_data.volume_groups:
        return
    _add_and_bookmark_heading(
        document, t("doc.headings.volume_groups", lang), 2, toc_list, counters
    )
    document.add_paragraph(t("doc.descriptions.volume_groups", lang))
    for vg in sorted(infra_data.volume_groups, key=lambda v: v.display_name):
        _add_and_bookmark_heading(
            document,
            f"{t('doc.common.volume_group', lang)}: {vg.display_name}",
            3,
            toc_list,
            counters,
        )
        info_data = {
            t("doc.common.name", lang): vg.display_name,
            t("doc.common.state", lang): vg.lifecycle_state,
            t("doc.common.ad", lang): vg.availability_domain,
            t("doc.common.member_count", lang): str(len(vg.members)),
        }
        _create_titled_key_value_table(
            document, t("doc.headings.general_info", lang), info_data
        )
        if vg.members:
            headers = [t("doc.headers.group_members", lang)]
            members_table = document.add_table(
                rows=1, cols=len(headers), style="Table Grid"
            )
            _style_table_headers(members_table, headers)
            for member_name in vg.members:
                members_table.add_row().cells[0].text = member_name
            document.add_paragraph()
        else:
            document.add_paragraph(t("doc.messages.no_vg_members", lang))
        val = vg.validation
        validation_data = {
            t("doc.common.backup_policy", lang): val.policy_name,
            t("doc.common.cross_region_replication", lang): t("doc.common.enabled", lang)
            if val.is_cross_region_replication_enabled
            else t("doc.common.disabled", lang),
            t("doc.common.replication_target", lang): val.cross_region_target,
        }
        _create_titled_key_value_table(
            document, t("doc.headings.data_protection_validation", lang), validation_data
        )


def _add_vcn_details_section(
    document: Document,
    infra_data: InfrastructureData,
    toc_list: list,
    counters: Dict[int, int],
    lang: str,
):
    """Adds the detailed VCN topology section, including subnets, SLs, RTs, NSGs, and LPGs."""
    sl_to_hosts, rt_to_hosts, nsg_to_hosts = {}, {}, {}
    for instance in infra_data.instances:
        for sl in instance.security_lists:
            sl_to_hosts.setdefault(sl.name, []).append(instance.host_name)
        if instance.route_table:
            rt_to_hosts.setdefault(instance.route_table.name, []).append(
                instance.host_name
            )
        for nsg in instance.network_security_groups:
            nsg_to_hosts.setdefault(nsg.name, []).append(instance.host_name)

    _add_and_bookmark_heading(
        document, t("doc.headings.vcn_topology", lang), 2, toc_list, counters
    )
    if not infra_data.vcns:
        document.add_paragraph(t("doc.messages.no_vcn_found", lang))
        return
    for vcn in infra_data.vcns:
        _add_and_bookmark_heading(
            document, f"VCN: {vcn.display_name}", 3, toc_list, counters
        )
        _create_titled_key_value_table(
            document,
            t("doc.headings.vcn_details", lang),
            {"CIDR Block": vcn.cidr_block},
            col_widths=(Inches(1.5), Inches(5.0)),
        )
        _add_and_bookmark_heading(
            document, t("doc.headings.subnets", lang), 4, toc_list, counters
        )
        if vcn.subnets:
            headers = [
                t("doc.headers.subnet_name", lang),
                t("doc.headers.cidr_block", lang),
            ]
            table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
            _style_table_headers(table, headers)
            for subnet in sorted(vcn.subnets, key=lambda s: s.display_name):
                cells = table.add_row().cells
                cells[0].text = subnet.display_name
                cells[1].text = subnet.cidr_block
        else:
            document.add_paragraph(t("doc.messages.no_subnet_found", lang))
        document.add_paragraph()

        _add_and_bookmark_heading(
            document, t("doc.headings.security_lists", lang), 4, toc_list, counters
        )
        if vcn.security_lists:
            for sl in sorted(vcn.security_lists, key=lambda s: s.name):
                _add_network_resource_details(
                    document,
                    t("doc.headings.sl_rules", lang),
                    sl.name,
                    sl.rules,
                    sl_to_hosts.get(sl.name),
                    "security",
                    lang,
                )
        else:
            document.add_paragraph(t("doc.messages.no_sl_found", lang))

        _add_and_bookmark_heading(
            document, t("doc.headings.route_tables", lang), 4, toc_list, counters
        )
        if vcn.route_tables:
            for rt in sorted(vcn.route_tables, key=lambda r: r.name):
                _add_network_resource_details(
                    document,
                    t("doc.headings.rt_rules", lang),
                    rt.name,
                    rt.rules,
                    rt_to_hosts.get(rt.name),
                    "route",
                    lang,
                )
        else:
            document.add_paragraph(t("doc.messages.no_rt_found", lang))

        _add_and_bookmark_heading(
            document, t("doc.headings.nsgs", lang), 4, toc_list, counters
        )
        if vcn.network_security_groups:
            for nsg in sorted(vcn.network_security_groups, key=lambda n: n.name):
                _add_network_resource_details(
                    document,
                    t("doc.headings.nsg_rules", lang),
                    nsg.name,
                    nsg.rules,
                    nsg_to_hosts.get(nsg.name),
                    "security",
                    lang,
                )
        else:
            document.add_paragraph(t("doc.messages.no_nsg_found", lang))

        _add_and_bookmark_heading(
            document, t("doc.headings.lpgs", lang), 4, toc_list, counters
        )
        if vcn.lpgs:
            headers = [
                t("doc.common.name", lang),
                t("doc.headers.peering_status", lang),
                t("doc.headers.route_table", lang),
                t("doc.headers.advertised_cidr", lang),
                t("doc.headers.cross_tenancy", lang),
            ]
            table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
            _style_table_headers(table, headers)
            for lpg in vcn.lpgs:
                cells = table.add_row().cells
                cells[0].text = lpg.display_name
                cells[1].text = lpg.peering_status_details or lpg.peering_status
                cells[2].text = lpg.route_table_name
                cells[3].text = lpg.peer_advertised_cidr or "N/A"
                cells[4].text = (
                    t("doc.common.yes", lang)
                    if lpg.is_cross_tenancy_peering
                    else t("doc.common.no", lang)
                )
        else:
            document.add_paragraph(t("doc.messages.no_lpg_found", lang))
        document.add_paragraph()


def _add_kubernetes_section(
    document: Document,
    infra_data: InfrastructureData,
    toc_list: list,
    counters: Dict[int, int],
    lang: str,
):
    """Adds the OKE (Kubernetes) clusters section to the document."""
    if (
        not hasattr(infra_data, "kubernetes_clusters")
        or not infra_data.kubernetes_clusters
    ):
        return

    _add_and_bookmark_heading(
        document, t("doc.headings.oke_config", lang), 1, toc_list, counters
    )
    for cluster in sorted(infra_data.kubernetes_clusters, key=lambda c: c.name):
        _add_and_bookmark_heading(
            document,
            f"{t('doc.headings.oke_cluster', lang)}: {cluster.name}",
            2,
            toc_list,
            counters,
        )
        headers = [
            t("doc.headers.cluster_name", lang),
            t("doc.headers.k8s_version", lang),
            t("doc.headers.vcn_associated", lang),
            t("doc.headers.public_api_endpoint", lang),
            t("doc.headers.private_api_endpoint", lang),
        ]
        table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
        _style_table_headers(table, headers)
        cells = table.add_row().cells
        cells[0].text = cluster.name
        cells[1].text = cluster.kubernetes_version
        cells[2].text = cluster.vcn_name
        cells[3].text = cluster.public_api_endpoint
        cells[4].text = cluster.private_api_endpoint
        document.add_paragraph()

        _add_and_bookmark_heading(
            document, t("doc.headings.oke_nodepools", lang), 3, toc_list, counters
        )
        if cluster.node_pools:
            headers = [
                t("doc.headers.nodepool_name", lang),
                t("doc.headers.shape", lang),
                t("doc.headers.ocpu", lang),
                t("doc.common.memory_gb", lang),
                t("doc.common.disk_gb", lang),
                t("doc.headers.os", lang),
                t("doc.headers.node_count", lang),
                t("doc.headers.subnet", lang),
            ]
            table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
            _style_table_headers(table, headers)
            for np in sorted(cluster.node_pools, key=lambda p: p.name):
                cells = table.add_row().cells
                cells[0].text = np.name
                cells[1].text = np.shape
                cells[2].text = str(np.ocpus)
                cells[3].text = str(np.memory_in_gbs)
                cells[4].text = str(np.boot_volume_size_in_gbs)
                cells[5].text = np.os_image
                cells[6].text = str(np.node_count)
                cells[7].text = np.subnet_name
            document.add_paragraph()
        else:
            document.add_paragraph(t("doc.messages.no_nodepool_found", lang))


def _add_load_balancers_section(
    document: Document,
    infra_data: InfrastructureData,
    toc_list: list,
    counters: Dict[int, int],
    lang: str,
):
    """Adds the Load Balancers (LBaaS) section to the document."""
    if not hasattr(infra_data, "load_balancers") or not infra_data.load_balancers:
        return
    _add_and_bookmark_heading(
        document, t("doc.headings.load_balancers", lang), 2, toc_list, counters
    )
    for lb in infra_data.load_balancers:
        _add_and_bookmark_heading(
            document,
            f"{t('doc.headings.lb_details', lang)}: {lb.display_name}",
            3,
            toc_list,
            counters,
        )
        _render_single_load_balancer(document, lb, infra_data, toc_list, counters, lang, base_level=3)



# Row background/text colors for VPN tunnel status.
# DOWN = amber (attention); UP = subtle green (ok); PROVISIONING = light orange.
_TUNNEL_STATUS_ROW_COLORS: Dict[str, Optional[str]] = {
    "DOWN":         "FEF3C7",   # amber-50: subtle yellow — attention without alarm
    "PROVISIONING": "FFF7ED",   # amber-25: very light
    "AVAILABLE":    None,       # no color — active/ok
    "UP":           "F0FDF4",   # green-50: very subtle green
}
_TUNNEL_STATUS_TEXT_COLORS: Dict[str, str] = {
    "DOWN":         "92400E",   # amber-800 brown — clearly readable
    "PROVISIONING": "C2410C",   # orange-700
    "UP":           "166534",   # green-800
}


def _style_row_by_tunnel_status(row, status_raw: str) -> None:
    """
    Applies subtle background/text color to a row based on VPN tunnel status.
    DOWN = amber-yellow (attention); UP = very light green (ok).
    """
    key        = (status_raw or "").upper()
    bg_color   = _TUNNEL_STATUS_ROW_COLORS.get(key)
    text_color = _TUNNEL_STATUS_TEXT_COLORS.get(key)
    if not bg_color and not text_color:
        return
    for cell in row.cells:
        if bg_color:
            _shade_cell(cell, bg_color)
        for para in cell.paragraphs:
            for run in para.runs:
                if text_color:
                    run.font.color.rgb = RGBColor.from_string(text_color)
            if para.text and not para.runs:
                text_val = para.text
                para.clear()
                run = para.add_run(text_val)
                if text_color:
                    run.font.color.rgb = RGBColor.from_string(text_color)


def _add_connectivity_section(
    document: Document,
    infra_data: InfrastructureData,
    toc_list: list,
    counters: Dict[int, int],
    lang: str,
):
    """Adds the external connectivity section — DRGs first, then CPEs and VPN tunnels."""
    _add_and_bookmark_heading(
        document, t("doc.headings.connectivity", lang), 2, toc_list, counters
    )

    # ── DRGs ────────────────────────────────────────────────────────────────────
    _add_and_bookmark_heading(
        document, t("doc.headings.drgs", lang), 3, toc_list, counters
    )
    if not infra_data.drgs:
        document.add_paragraph(t("doc.messages.no_drg_found", lang))
    else:
        # Build VCN id→name map for attachment enrichment
        vcn_name_map = {vcn.id: vcn.display_name for vcn in (infra_data.vcns or [])}
        for drg in infra_data.drgs:
            _add_and_bookmark_heading(
                document,
                f"{t('doc.headings.drg_details', lang)}: {drg.display_name}",
                4,
                toc_list,
                counters,
            )
            # Attachments
            _add_and_bookmark_heading(
                document, t("doc.headings.drg_attachments", lang), 5, toc_list, counters
            )
            if drg.attachments:
                headers = [
                    t("doc.headers.drg_attachment_name", lang),
                    t("doc.headers.drg_attachment_type", lang),
                    "VCN",
                    t("doc.headers.drg_route_table", lang),
                ]
                table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
                _style_table_headers(table, headers)
                for attachment in drg.attachments:
                    vcn_name = ""
                    if (attachment.network_type or "").upper() == "VCN":
                        vcn_name = vcn_name_map.get(attachment.network_id or "", "—")
                    cells = table.add_row().cells
                    cells[0].text = attachment.display_name
                    cells[1].text = attachment.network_type
                    cells[2].text = vcn_name
                    cells[3].text = attachment.route_table_name or "N/A"
                document.add_paragraph()
            else:
                document.add_paragraph(t("doc.messages.no_drg_attachment_found", lang))

            # RPCs
            _add_and_bookmark_heading(
                document, t("doc.headings.rpcs", lang), 5, toc_list, counters
            )
            if drg.rpcs:
                headers = [
                    t("doc.headers.rpc_name", lang),
                    t("doc.common.status", lang),
                    t("doc.headers.rpc_peering_status", lang),
                ]
                table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
                _style_table_headers(table, headers)
                for rpc in drg.rpcs:
                    status_text = rpc.peering_status_details or rpc.peering_status
                    if rpc.peering_status == "NEW":
                        status_text = "New (not peered)"
                    cells = table.add_row().cells
                    cells[0].text = rpc.display_name
                    cells[1].text = rpc.lifecycle_state
                    cells[2].text = status_text
                document.add_paragraph()
            else:
                document.add_paragraph(t("doc.messages.no_rpc_found", lang))

    # ── CPEs ────────────────────────────────────────────────────────────────────
    _add_and_bookmark_heading(
        document, t("doc.headings.cpes", lang), 3, toc_list, counters
    )
    if not infra_data.cpes:
        document.add_paragraph(t("doc.messages.no_cpe_found", lang))
    else:
        headers = [
            t("doc.headers.cpe_name", lang),
            t("doc.headers.cpe_ip", lang),
            t("doc.headers.cpe_vendor", lang),
        ]
        table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
        _style_table_headers(table, headers)
        for cpe in infra_data.cpes:
            cells = table.add_row().cells
            cells[0].text = cpe.display_name
            cells[1].text = cpe.ip_address
            cells[2].text = cpe.vendor or "N/A"
    document.add_paragraph()

    # ── VPN / IPSec ──────────────────────────────────────────────────────────────
    _add_and_bookmark_heading(
        document, t("doc.headings.vpn_connections", lang), 3, toc_list, counters
    )
    if not infra_data.ipsec_connections:
        document.add_paragraph(t("doc.messages.no_vpn_found", lang))
    else:
        cpe_map = {cpe.id: cpe.display_name for cpe in infra_data.cpes}
        drg_map = {drg.id: drg.display_name for drg in infra_data.drgs}

        # Summary table
        _add_and_bookmark_heading(
            document, t("doc.headings.vpn_summary", lang), 4, toc_list, counters
        )
        headers = [
            t("doc.headers.vpn_connection_name", lang),
            t("doc.common.status", lang),
            t("doc.headers.vpn_cpe_associated", lang),
            t("doc.headers.vpn_drg_associated", lang),
            t("doc.headers.vpn_routing", lang),
            t("doc.headers.vpn_tunnels", lang),
        ]
        tbl_summary = document.add_table(rows=1, cols=len(headers), style="Table Grid")
        _style_table_headers(tbl_summary, headers)
        for ipsec in infra_data.ipsec_connections:
            routing_type = (
                "BGP" if any(tn.routing_type == "BGP" for tn in ipsec.tunnels) else "STATIC"
            )
            routing_display = (
                f"STATIC ({len(ipsec.static_routes)} rotas)"
                if routing_type == "STATIC"
                else "BGP"
            )
            cells = tbl_summary.add_row().cells
            cells[0].text = ipsec.display_name
            cells[1].text = ipsec.status
            cells[2].text = cpe_map.get(ipsec.cpe_id, ipsec.cpe_id)
            cells[3].text = drg_map.get(ipsec.drg_id, ipsec.drg_id)
            cells[4].text = routing_display
            cells[5].text = str(len(ipsec.tunnels))
            # Color the summary row by connection status (UP/DOWN/PROVISIONING)
            _style_row_by_tunnel_status(tbl_summary.rows[-1], ipsec.status)
        document.add_paragraph()

        # Detailed tunnel tables
        _add_and_bookmark_heading(
            document, t("doc.headings.vpn_details", lang), 4, toc_list, counters
        )
        for ipsec in infra_data.ipsec_connections:
            document.add_paragraph(
                f"{t('doc.headings.vpn_connection_config', lang)}: {ipsec.display_name}",
                style="Heading 5",
            )
            if not ipsec.tunnels:
                document.add_paragraph(t("doc.messages.no_tunnel_found", lang))
                continue
            for tunnel in ipsec.tunnels:
                # ── Tunnel header paragraph with bold name + status badge ──────
                p = document.add_paragraph()
                run_name = p.add_run(
                    f"{t('doc.headings.vpn_tunnel_details', lang)}: {tunnel.display_name}  "
                )
                run_name.bold = True
                # Status inline badge: DOWN → amber text, UP → green text
                status_upper = (tunnel.status or "").upper()
                run_status = p.add_run(f"[{tunnel.status}]")
                run_status.bold = True
                if status_upper == "DOWN":
                    run_status.font.color.rgb = RGBColor.from_string("92400E")
                elif status_upper == "UP":
                    run_status.font.color.rgb = RGBColor.from_string("166534")
                else:
                    run_status.font.color.rgb = RGBColor.from_string("C2410C")

                tunnel_info = {
                    t("doc.headers.vpn_oracle_ip", lang): tunnel.vpn_oracle_ip or "N/A",
                    t("doc.headers.vpn_cpe_ip", lang):    tunnel.cpe_ip or "N/A",
                    t("doc.common.routing", lang):         tunnel.routing_type,
                    t("doc.headers.vpn_ike_version", lang): tunnel.ike_version,
                }
                _create_titled_key_value_table(
                    document, t("doc.headings.vpn_tunnel_info", lang), tunnel_info
                )
                # Apply subtle row background to all data rows based on tunnel status
                tbl_tunnel = document.tables[-1]
                for row in tbl_tunnel.rows[1:]:
                    _style_row_by_tunnel_status(row, tunnel.status)

                if tunnel.routing_type == "BGP" and tunnel.bgp_session_info:
                    bgp = tunnel.bgp_session_info
                    bgp_info_data = {
                        t("doc.headers.vpn_bgp_oracle_asn", lang):    bgp.oracle_bgp_asn,
                        t("doc.headers.vpn_bgp_customer_asn", lang):   bgp.customer_bgp_asn,
                        t("doc.headers.vpn_bgp_oracle_ip", lang):      bgp.oracle_interface_ip,
                        t("doc.headers.vpn_bgp_customer_ip", lang):    bgp.customer_interface_ip,
                    }
                    _create_titled_key_value_table(
                        document, t("doc.headings.vpn_bgp_details", lang), bgp_info_data
                    )

                p1 = tunnel.phase_one_details
                p1_info = {
                    t("doc.headers.vpn_p1_auth", lang):       p1.authentication_algorithm,
                    t("doc.headers.vpn_p1_encryption", lang): p1.encryption_algorithm,
                    t("doc.headers.vpn_p1_dh_group", lang):   p1.dh_group,
                    t("doc.headers.vpn_p1_lifetime", lang):   str(p1.lifetime_in_seconds),
                }
                _create_titled_key_value_table(
                    document, t("doc.headings.vpn_phase1", lang), p1_info
                )

                p2 = tunnel.phase_two_details
                p2_info = {
                    t("doc.headers.vpn_p2_auth", lang):       p2.authentication_algorithm or "N/A",
                    t("doc.headers.vpn_p2_encryption", lang): p2.encryption_algorithm,
                    t("doc.headers.vpn_p2_lifetime", lang):   str(p2.lifetime_in_seconds),
                }
                _create_titled_key_value_table(
                    document, t("doc.headings.vpn_phase2", lang), p2_info
                )
                if (
                    tunnel.validation_status == "Fora da recomendação Oracle"
                    and tunnel.validation_details
                ):
                    p = document.add_paragraph()
                    p.add_run(t("doc.headings.vpn_config_note", lang)).bold = True
                    p.add_run(t("doc.descriptions.vpn_config_warning", lang))
                    url_p = document.add_paragraph(tunnel.validation_details)
                    url_p.style = "Intense Quote"
                document.add_paragraph()


def _add_responsible_section(
    document: Document,
    toc_list: list,
    counters: Dict[int, int],
    responsible_name: str,
    lang: str,
):
    """Adds the final 'Responsible' signature section to the document."""
    _add_and_bookmark_heading(
        document, t("doc.headings.responsible", lang), 1, toc_list, counters
    )
    document.add_paragraph(t("doc.headings.responsible_desc", lang))
    headers = [t("doc.headers.responsible", lang), t("doc.headers.date", lang)]
    table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
    _style_table_headers(table, headers)
    cells = table.add_row().cells
    cells[0].text = responsible_name
    cells[1].text = datetime.now().strftime("%d/%m/%Y")
    document.add_paragraph()


# --- Document Body Helpers ---
def _add_network_resource_details(
    document: Document,
    resource_title: str,
    resource_name: str,
    rules: List[Any],
    associated_hosts: Optional[List[str]],
    rule_type: str,
    lang: str,
):
    """Adds a standardized section for a network resource (SL, NSG, RT) and its rules."""
    p = document.add_paragraph()
    p.add_run(f"{resource_title}: {resource_name}").bold = True

    if associated_hosts:
        p_hosts = document.add_paragraph()
        p_hosts.add_run(t("doc.common.associated_hosts", lang)).bold = True
        p_hosts.add_run(", ".join(sorted(associated_hosts)))

    if rule_type == "security":
        headers = [
            t("doc.headers.direction", lang),
            t("doc.headers.protocol", lang),
            t("doc.headers.ports", lang),
            t("doc.headers.source_destination", lang),
            t("doc.headers.description", lang),
        ]
    else:  # route
        headers = [
            t("doc.headers.destination", lang),
            t("doc.headers.target", lang),
            t("doc.headers.description", lang),
        ]

    if rules:
        table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
        _style_table_headers(table, headers)
        for rule in rules:
            cells = table.add_row().cells
            if rule_type == "security":
                cells[0].text = rule.direction
                cells[1].text = rule.protocol
                cells[2].text = rule.ports or t("doc.common.all", lang)
                cells[3].text = rule.source_or_destination or "N/A"
                cells[4].text = rule.description or ""
            else:  # route
                target_str = rule.target
                # If the backend could not resolve the display name, the raw OCID is sent — label it generically.
                if target_str and target_str.startswith("ocid1."):
                    target_lower = target_str.lower()
                    if "internetgateway" in target_lower:
                        target_str = "Internet Gateway"
                    elif "natgateway" in target_lower:
                        target_str = "NAT Gateway"
                    elif "servicegateway" in target_lower:
                        target_str = "Service Gateway"
                    elif "drg" in target_lower:
                        target_str = "Dynamic Routing Gateway (DRG)"
                    elif "localpeeringgateway" in target_lower:
                        target_str = "Local Peering Gateway (LPG)"
                    elif "privateip" in target_lower:
                        target_str = "Private IP"

                cells[0].text = rule.destination
                cells[1].text = target_str
                cells[2].text = rule.description or ""
    else:
        document.add_paragraph(f"{t('doc.messages.no_rules_for', lang)} {resource_title}.")

    document.add_paragraph()


def _insert_image_sections(
    document,
    sections: list,
    position: str,
    toc_list: list,
    counters: dict,
    lang: str,
) -> None:
    """Insert user-defined image sections into the document.

    Each section: level-1 heading → optional text_above → images → optional text_below.
    Only sections whose position matches the given argument are inserted.
    """
    for sec in sections:
        if sec.get("position") != position:
            continue
        images = sec.get("images") or []
        if not images:
            continue
        section_name = sec.get("name", "Anexo")
        _add_and_bookmark_heading(document, section_name, 1, toc_list, counters)

        text_above = (sec.get("text_above") or "").strip()
        if text_above:
            p = document.add_paragraph(text_above)
            p.style.font.italic = True

        for image_bytes in images:
            try:
                document.add_picture(BytesIO(image_bytes), width=Inches(6.0))
                document.add_paragraph()
            except Exception as e:
                document.add_paragraph(f"{t('doc.messages.insert_image_error', lang)}: {e}")

        text_below = (sec.get("text_below") or "").strip()
        if text_below:
            p = document.add_paragraph(text_below)
            p.style.font.italic = True




def _apply_letterhead(
    document,
    header_bytes: Optional[bytes],
    footer_bytes: Optional[bytes],
) -> None:
    """Stamp header/footer images using the full-page-width bleed technique.

    Replicates exactly what the CCM template does:

    1. header_distance = footer_distance = 0  →  header area starts at y=0 (top of page).
    2. w:ind hanging = left_margin_twips  →  negative indent pulls the image LEFT
       by exactly the left margin, so the image starts at the physical left page edge.
    3. Image width = section.page_width (full physical page width, not text width).

    This combination produces a header/footer that bleeds edge-to-edge with no
    lateral gaps regardless of the document's margin settings.

    Footer layout
    -------------
    Paragraph 1 (hanging indent): footer image at full page width.
    Paragraph 2 (right-aligned):  page number field below the image.
    """
    from docx.oxml.ns import qn as _qn
    from docx.oxml.shared import OxmlElement as _OE

    def _fld_char(kind: str):
        el = _OE("w:fldChar")
        el.set(_qn("w:fldCharType"), kind)
        return el

    def _instr(text: str):
        el = _OE("w:instrText")
        el.set(_qn("xml:space"), "preserve")
        el.text = text
        return el

    def _add_page_num(para):
        """Insert a PAGE field run into *para*."""
        run = para.add_run()
        run._r.append(_fld_char("begin"))
        run._r.append(_instr(" PAGE "))
        run._r.append(_fld_char("separate"))
        run._r.append(_fld_char("end"))
        run.font.name = "Gotham Light"
        run.font.size = Pt(9)

    def _set_bleed_para(para, left_twips: int):
        """Apply the hanging-indent bleed technique to *para*.

        Removes all left/right/firstLine ind overrides and sets ONLY
        w:hanging = left_margin_twips.  This mirrors the pPr of the template
        header/footer paragraphs exactly.
        """
        pPr = para._p.get_or_add_pPr()
        # Remove any existing ind and spacing elements
        for tag in (_qn("w:ind"), _qn("w:spacing")):
            for el in pPr.findall(tag):
                pPr.remove(el)
        # Hanging indent = left margin → image starts at left page edge
        ind = _OE("w:ind")
        ind.set(_qn("w:hanging"), str(left_twips))
        pPr.append(ind)
        # Zero all spacing so there is no gap above/below the image
        spc = _OE("w:spacing")
        spc.set(_qn("w:before"), "0")
        spc.set(_qn("w:after"), "0")
        pPr.append(spc)

    for section in document.sections:
        # ── Compute per-section values ────────────────────────────────────────
        # Left margin in twips (20 twips = 1 pt = 1/72 inch;  1 EMU = 914400 /inch)
        left_twips = int(round(section.left_margin / 914400 * 1440))
        # Full physical page width in EMU — this is what python-docx add_picture expects
        page_width_emu = int(section.page_width)

        # Set header and footer distances to 0 so they start at the page edge.
        # This is the critical difference from the default (0.5 inch).
        section.header_distance = Pt(0)
        section.footer_distance = Pt(0)

        # ── Header ────────────────────────────────────────────────────────────
        if header_bytes:
            section.header.is_linked_to_previous = False
            hdr_para = (
                section.header.paragraphs[0]
                if section.header.paragraphs
                else section.header.add_paragraph()
            )
            hdr_para.clear()
            hdr_para.alignment = None  # Do NOT set explicit alignment — let hanging handle layout
            _set_bleed_para(hdr_para, left_twips)
            try:
                hdr_para.add_run().add_picture(BytesIO(header_bytes), width=page_width_emu)
            except Exception:
                hdr_para.add_run("[Cabeçalho]")

        # ── Footer ────────────────────────────────────────────────────────────
        section.footer.is_linked_to_previous = False
        for p in section.footer.paragraphs:
            p.clear()
        ftr_para = (
            section.footer.paragraphs[0]
            if section.footer.paragraphs
            else section.footer.add_paragraph()
        )
        ftr_para.clear()
        _set_bleed_para(ftr_para, left_twips)

        if footer_bytes:
            ftr_para.alignment = None  # Hanging indent handles lateral position
            try:
                ftr_para.add_run().add_picture(BytesIO(footer_bytes), width=page_width_emu)
            except Exception:
                ftr_para.add_run("[Rodapé]")
            # Page number in a separate paragraph below the footer image
            num_para = section.footer.add_paragraph()
            _set_bleed_para(num_para, left_twips)
            num_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _add_page_num(num_para)
        else:
            # No footer image — just a right-aligned page number
            ftr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _add_page_num(ftr_para)


# --- Document Generation ---
def generate_documentation(
    doc_type: str,
    infra_data: InfrastructureData,
    responsible_name: str,
    image_sections: Optional[List[dict]] = None,
    lang: str = "pt",
    compartment_name: Optional[str] = None,
    letterhead: Optional[dict] = None,
) -> str:
    """Main function to generate a .docx file.

    image_sections: list of dicts {name, position ("start"|"end"), images: List[bytes]}
    compartment_name: explicit compartment name passed from the API request; used as the
                      primary source for the "Cliente" field so it is always populated
                      regardless of which resources exist in the collected data.
    letterhead: optional dict {enabled, header, footer, cover} — header/footer images are
                stamped into every section using the full-page-width bleed technique.
    """
    document = Document()
    # ── Global font: Gotham Light for body, Gotham Medium for headings/title ──
    _ns = document.styles["Normal"]
    _ns.font.name = "Gotham Light"
    _ns.font.size = Pt(11)
    for _hl in range(1, 10):
        _hs = f"Heading {_hl}"
        if _hs in document.styles:
            document.styles[_hs].font.name = "Gotham Medium"
    if "Title" in document.styles:
        document.styles["Title"].font.name = "Gotham Medium"
    _define_toc_styles(document)

    # Prefer the explicit compartment_name from the API request (always populated).
    # Fall back to deriving it from collected data for backwards compatibility.
    if compartment_name and compartment_name.strip() not in ("", "N/A"):
        client_name = compartment_name.strip()
    elif infra_data.instances:
        client_name = infra_data.instances[0].compartment_name.replace("SERVERS-", "")
    elif infra_data.kubernetes_clusters:
        client_name = "Compartimento_OKE"
    elif hasattr(infra_data, "waf_policies") and infra_data.waf_policies:
        client_name = infra_data.waf_policies[0].compartment_name
    elif infra_data.vcns:
        client_name = "Compartimento"
    else:
        client_name = "Desconhecido"

    safe_client_name = re.sub(r'[\\/*?:"<>|]', "", client_name)
    doc_type_map = {
        "full_infra": (t("doc.type.full_infra", lang), t("doc.identifier.full_infra", lang)),
        "new_host": (t("doc.type.new_host", lang), t("doc.identifier.new_host", lang)),
        "kubernetes": (t("doc.type.kubernetes", lang), t("doc.identifier.kubernetes", lang)),
        "waf_report": (t("doc.type.waf_report", lang), t("doc.identifier.waf_report", lang)),
    }
    doc_title_text, doc_identifier = doc_type_map.get(
        doc_type, (t("doc.type.default", lang), t("doc.identifier.default", lang))
    )

    headings_for_toc: List[Tuple[str, str, int]] = []
    numbering_counters: Dict[int, int] = {i: 0 for i in range(1, 6)}

    # ── Cover page ────────────────────────────────────────────────────────────
    cover_bytes = (letterhead or {}).get("cover") if letterhead else None
    if cover_bytes:
        cover_img_p = document.add_paragraph()
        cover_img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt = cover_img_p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after  = Pt(0)
        try:
            cover_img_p.add_run().add_picture(BytesIO(cover_bytes), width=Inches(6.27))
        except Exception:
            cover_img_p.add_run("[Imagem de Capa]")
        document.add_paragraph()
        title_p = document.add_paragraph(doc_title_text, style="Title")
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_p = document.add_paragraph(client_name)
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p = document.add_paragraph(
            f"{t('doc.common.generation_date', lang)}: {datetime.now().strftime('%d/%m/%Y')}"
        )
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        title_p = document.add_paragraph(doc_title_text, style="Title")
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(
            f"{t('doc.common.client', lang)}: {client_name}\n{t('doc.common.generation_date', lang)}: {datetime.now().strftime('%d/%m/%Y')}"
        )
    document.add_page_break()
    toc_placeholder = document.add_paragraph(t("doc.common.toc", lang), style="Heading 1")
    document.add_page_break()

    # Image sections positioned at "start" are injected before infrastructure content.
    _insert_image_sections(document, image_sections or [], "start", headings_for_toc, numbering_counters, lang)

    if doc_type == "kubernetes":
        cluster_vcn_ids = {cluster.vcn_id for cluster in infra_data.kubernetes_clusters}
        relevant_vcns = [vcn for vcn in infra_data.vcns if vcn.id in cluster_vcn_ids]
        filtered_infra_data = infra_data.copy(deep=True)
        filtered_infra_data.vcns = relevant_vcns
        filtered_infra_data.instances = []
        filtered_infra_data.drgs = []
        filtered_infra_data.cpes = []
        filtered_infra_data.ipsec_connections = []
        filtered_infra_data.load_balancers = []
        filtered_infra_data.volume_groups = []
        _add_kubernetes_section(document, filtered_infra_data, headings_for_toc, numbering_counters, lang)
        _add_vcn_details_section(document, filtered_infra_data, headings_for_toc, numbering_counters, lang)

    elif doc_type == "waf_report":
        # 1. Filter active WAF policies only.
        active_waf_policies = []
        if hasattr(infra_data, "waf_policies") and infra_data.waf_policies:
            active_waf_policies = [p for p in infra_data.waf_policies if getattr(p, "lifecycle_state", "").upper() == "ACTIVE"]

        target_lb_names = set()
        target_subnet_ids = set()
        target_subnet_names = set()
        target_vcn_names = set()
        target_lbs = []

        # 2. Extract LBs and network references tied to each firewall.
        # Iterate over all policy integrations (one policy may bind multiple LBs).
        for policy in active_waf_policies:
            all_integrations = getattr(policy, "integrations", []) or []
            if not all_integrations and getattr(policy, "integration", None):
                all_integrations = [policy.integration]

            for intg in all_integrations:
                lb = getattr(intg, "load_balancer", None)
                if lb:
                    lb_name = getattr(lb, "display_name", "Unknown_LB")
                    if lb_name not in target_lb_names:
                        target_lb_names.add(lb_name)
                        target_lbs.append(lb)
                        for subnet_id in getattr(lb, "subnet_ids", []):
                            target_subnet_ids.add(subnet_id)

            # Fall back to network_infrastructure when direct LB data is unavailable.
            net = getattr(policy, "network_infrastructure", None)
            if net:
                if getattr(net, "subnet_name", "N/A") != "N/A":
                    target_subnet_names.add(net.subnet_name)
                if getattr(net, "vcn_name", "N/A") != "N/A":
                    target_vcn_names.add(net.vcn_name)

        # 3. Filter VCNs to only those referenced by the WAF/LB scope.
        filtered_vcns = []
        if target_subnet_ids or target_subnet_names or target_vcn_names:
            for vcn in getattr(infra_data, "vcns", []):
                vcn_name = getattr(vcn, "display_name", getattr(vcn, "name", ""))
                vcn_has_target = False
                matching_subnets = []
                matching_sl_ids = set()
                matching_rt_ids = set()

                for sub in getattr(vcn, "subnets", []):
                    sub_id = getattr(sub, "id", None)
                    sub_name = getattr(sub, "display_name", getattr(sub, "name", None))
                    
                    if (sub_id and sub_id in target_subnet_ids) or (sub_name and sub_name in target_subnet_names) or (vcn_name in target_vcn_names):
                        vcn_has_target = True
                        matching_subnets.append(sub)
                        for sl_id in getattr(sub, "security_list_ids", []):
                            matching_sl_ids.add(sl_id)
                        if getattr(sub, "route_table_id", None):
                            matching_rt_ids.add(getattr(sub, "route_table_id"))

                if not vcn_has_target and vcn_name in target_vcn_names:
                    vcn_has_target = True
                    matching_subnets = getattr(vcn, "subnets", [])

                if vcn_has_target:
                    filtered_vcn = vcn.copy(deep=True)
                    filtered_vcn.subnets = matching_subnets
                    if getattr(filtered_vcn, "security_lists", None) and matching_sl_ids:
                        filtered_vcn.security_lists = [sl for sl in filtered_vcn.security_lists if getattr(sl, "id", None) in matching_sl_ids]
                    if getattr(filtered_vcn, "route_tables", None) and matching_rt_ids:
                        filtered_vcn.route_tables = [rt for rt in filtered_vcn.route_tables if getattr(rt, "id", None) in matching_rt_ids]
                    filtered_vcn.network_security_groups = []
                    filtered_vcn.lpgs = []
                    filtered_vcns.append(filtered_vcn)

        # 4. Build an isolated payload so document section numbering is clean.
        filtered_infra_data = infra_data.copy(deep=True)
        filtered_infra_data.instances = []
        filtered_infra_data.drgs = []
        filtered_infra_data.cpes = []
        filtered_infra_data.ipsec_connections = []
        filtered_infra_data.volume_groups = []
        filtered_infra_data.kubernetes_clusters = []
        filtered_infra_data.load_balancers = target_lbs
        filtered_infra_data.vcns = filtered_vcns
        filtered_infra_data.waf_policies = active_waf_policies
        # Preserve certificates — the deep copy carries them from infra_data.
        if not getattr(filtered_infra_data, "certificates", None):
            filtered_infra_data.certificates = getattr(infra_data, "certificates", []) or []

        # 5. Render sections: VCN → LB → WAF → Certificates.
        if target_lbs or filtered_vcns:
            _add_and_bookmark_heading(document, t("doc.headings.infra_config", lang), 1, headings_for_toc, numbering_counters)
            if filtered_vcns:
                _add_vcn_details_section(document, filtered_infra_data, headings_for_toc, numbering_counters, lang)
            else:
                document.add_paragraph(t("doc.messages.no_network_found", lang))

            if target_lbs:
                _add_load_balancers_section(document, filtered_infra_data, headings_for_toc, numbering_counters, lang)
            else:
                document.add_paragraph(t("doc.messages.no_lb_association", lang))

        _add_waf_report_section(document, filtered_infra_data, headings_for_toc, numbering_counters, lang)

        # Certificates section — shows ACTIVE and PENDING_DELETION states.
        waf_active_certs = [
            c for c in (getattr(filtered_infra_data, "certificates", []) or [])
            if isinstance(c, dict)
            and c.get("lifecycle_state", "").upper() in ("ACTIVE", "PENDING_DELETION")
        ]
        if waf_active_certs:
            _add_and_bookmark_heading(
                document,
                t("doc.headers.certificates", lang),
                1, headings_for_toc, numbering_counters,
            )
            _add_compartment_certificates_section(
                document, filtered_infra_data, headings_for_toc, numbering_counters, lang
            )
        
    elif doc_type == "new_host":
        _add_and_bookmark_heading(document, t("doc.headings.infra_config", lang), 1, headings_for_toc, numbering_counters)
        _add_and_bookmark_heading(document, t("doc.headings.compute_instances", lang), 2, headings_for_toc, numbering_counters)
        _add_instances_table(document, infra_data.instances, lang)
        _add_volume_and_backup_section(document, infra_data, headings_for_toc, numbering_counters, lang)
        if hasattr(infra_data, "volume_groups") and infra_data.volume_groups:
            _add_volume_groups_section(document, infra_data, headings_for_toc, numbering_counters, lang)
        
        _add_and_bookmark_heading(document, t("doc.headings.instance_connectivity", lang), 2, headings_for_toc, numbering_counters)
        sl_map, nsg_map, rt_map = {}, {}, {}
        for data in infra_data.instances:
            for sl in data.security_lists:
                entry = sl_map.setdefault(sl.name, {"rules": sl.rules, "hosts": []})
                if data.host_name not in entry["hosts"]: entry["hosts"].append(data.host_name)
            for nsg in data.network_security_groups:
                entry = nsg_map.setdefault(nsg.name, {"rules": nsg.rules, "hosts": []})
                if data.host_name not in entry["hosts"]: entry["hosts"].append(data.host_name)
            if data.route_table:
                entry = rt_map.setdefault(data.route_table.name, {"rules": data.route_table.rules, "hosts": []})
                if data.host_name not in entry["hosts"]: entry["hosts"].append(data.host_name)

        _add_and_bookmark_heading(document, t("doc.headings.security_lists", lang), 3, headings_for_toc, numbering_counters)
        for name, info in sorted(sl_map.items()):
            _add_network_resource_details(document, t("doc.headings.sl_rules", lang), name, info["rules"], info.get("hosts"), "security", lang)
            
        _add_and_bookmark_heading(document, t("doc.headings.nsgs", lang), 3, headings_for_toc, numbering_counters)
        for name, info in sorted(nsg_map.items()):
            _add_network_resource_details(document, t("doc.headings.nsg_rules", lang), name, info["rules"], info.get("hosts"), "security", lang)
            
        _add_and_bookmark_heading(document, t("doc.headings.route_tables", lang), 3, headings_for_toc, numbering_counters)
        for name, info in sorted(rt_map.items()):
            _add_network_resource_details(document, t("doc.headings.rt_rules", lang), name, info["rules"], info.get("hosts"), "route", lang)
            
    else:  # 'full_infra'
        _add_and_bookmark_heading(document, t("doc.headings.infra_config", lang), 1, headings_for_toc, numbering_counters)
        # 1. Compute instances
        _add_and_bookmark_heading(document, t("doc.headings.compute_instances", lang), 2, headings_for_toc, numbering_counters)
        _add_instances_table(document, infra_data.instances, lang)
        # 2. Storage and volumes
        _add_volume_and_backup_section(document, infra_data, headings_for_toc, numbering_counters, lang)
        if hasattr(infra_data, "volume_groups") and infra_data.volume_groups:
            _add_volume_groups_section(document, infra_data, headings_for_toc, numbering_counters, lang)
        # 3. VCN topology
        _add_vcn_details_section(document, infra_data, headings_for_toc, numbering_counters, lang)
        # 4. Load balancers
        if hasattr(infra_data, "load_balancers") and infra_data.load_balancers:
            _add_load_balancers_section(document, infra_data, headings_for_toc, numbering_counters, lang)
        # 5. Certificates
        active_certs = [
            c for c in (getattr(infra_data, "certificates", []) or [])
            if isinstance(c, dict)
            and c.get("lifecycle_state", "").upper() in ("ACTIVE", "PENDING_DELETION")
        ]
        if active_certs:
            _add_and_bookmark_heading(
                document,
                t("doc.headers.certificates", lang),
                1, headings_for_toc, numbering_counters,
            )
            _add_compartment_certificates_section(
                document, infra_data, headings_for_toc, numbering_counters, lang
            )
        # 6. WAF policies
        _add_waf_report_section(
            document, infra_data, headings_for_toc, numbering_counters, lang
        )
        # 7. External connectivity (DRG → CPE → VPN)
        _add_connectivity_section(document, infra_data, headings_for_toc, numbering_counters, lang)
        # 8. Kubernetes (OKE) — only rendered when clusters are present.
        _add_kubernetes_section(document, infra_data, headings_for_toc, numbering_counters, lang)

    # Image sections positioned at "end" are injected after infrastructure content.
    _insert_image_sections(document, image_sections or [], "end", headings_for_toc, numbering_counters, lang)

    _add_responsible_section(document, headings_for_toc, numbering_counters, responsible_name, lang)

    # Finalize: generate the Table of Contents field and save to disk.
    for text, bookmark, level in reversed(headings_for_toc):
        style_name = f"TOC {level}" if f"TOC {level}" in document.styles else "TOC 1"
        p = document.add_paragraph(style=style_name)
        _add_internal_hyperlink(p, text, bookmark)
        toc_placeholder._p.addnext(p._p)

    # Apply header/footer images using full-page-width bleed technique.
    if letterhead and letterhead.get("enabled") and (letterhead.get("header") or letterhead.get("footer")):
        _apply_letterhead(document, header_bytes=letterhead.get("header"), footer_bytes=letterhead.get("footer"))

    output_dir = "generated_docs"
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_name = f"Doc_{doc_identifier}_{safe_client_name}_{timestamp}.docx"
    output_path = os.path.join(output_dir, file_name)

    document.save(output_path)
    return output_path


def _add_waf_report_section(
    document: Document,
    infra_data: InfrastructureData,
    toc_list: list,
    counters: Dict[int, int],
    lang: str,
):
    """Adds the WAF Policies details section to the document."""
    if not hasattr(infra_data, "waf_policies") or not infra_data.waf_policies:
        return

    _add_and_bookmark_heading(
        document, t("doc.headings.waf_config", lang), 1, toc_list, counters
    )

    for policy in sorted(infra_data.waf_policies, key=lambda p: p.display_name):
        _add_and_bookmark_heading(
            document,
            f"{t('doc.headings.waf_policy_name', lang)}: {policy.display_name}",
            2, toc_list, counters,
        )

        # 1. Overview section
        _add_and_bookmark_heading(document, t("doc.headings.overview", lang), 3, toc_list, counters)
        _create_titled_key_value_table(document, t("doc.headings.general_info", lang), {
            t("doc.common.name",        lang): policy.display_name,
            "OCID":                            policy.id,
            t("doc.common.compartment", lang): policy.compartment_name,
            t("doc.common.region",      lang): policy.region,
            t("doc.common.state",       lang): policy.lifecycle_state,
            t("doc.common.time_created",lang): policy.time_created,
        })

        # ── 2. Firewalls and bound Load Balancers (single consolidated table) ──────
        # One row per complete Firewall → LB binding, removing the duplication
        # that existed between the former separate firewall and LB sections.
        _add_and_bookmark_heading(document, t("doc.headings.waf_firewall", lang), 3, toc_list, counters)

        all_integrations = getattr(policy, "integrations", []) or []
        if not all_integrations and getattr(policy, "integration", None):
            all_integrations = [policy.integration]

        if not all_integrations:
            document.add_paragraph("Esta política WAF não possui Web Application Firewall associado.")
            continue

        # Binding table without OCIDs — already documented in LB General Info
        # and Policy Overview sections.
        # Columns: Firewall | Bound LB | IPs | State | Enforcement Point
        binding_headers = [
            t("doc.headers.firewall_name", lang),
            t("doc.headers.lb_name",       lang),
            t("doc.headers.lb_ip",         lang),
            t("doc.common.state",          lang),
            "Enforcement Point",
        ]
        binding_table = document.add_table(rows=1, cols=len(binding_headers), style="Table Grid")
        _style_table_headers(binding_table, binding_headers)

        for intg in all_integrations:
            fw     = getattr(intg, "firewall",      None)
            lb_obj = getattr(intg, "load_balancer", None)
            if not fw:
                continue

            full_lb = next(
                (lb for lb in (getattr(infra_data, "load_balancers", []) or [])
                 if lb_obj and getattr(lb, "display_name", "") == getattr(lb_obj, "display_name", "")),
                None,
            )
            lb_show = full_lb or lb_obj

            ip_strings = ", ".join(
                f"{ip.ip_address} ({t('doc.common.public', lang) if ip.is_public else t('doc.common.private', lang)})"
                for ip in (getattr(lb_show, "ip_addresses", []) or [])
            ) if lb_show else "N/A"

            lb_state = getattr(lb_show, "lifecycle_state", "N/A") if lb_show else "N/A"

            cells = binding_table.add_row().cells
            cells[0].text = getattr(fw,      "display_name", "N/A")
            cells[1].text = getattr(lb_show, "display_name", "N/A") if lb_show else "N/A"
            cells[2].text = ip_strings or "N/A"
            cells[3].text = lb_state
            cells[4].text = "Load Balancer"
            _style_row_by_state(binding_table.rows[-1], lb_state)

        document.add_paragraph()

        # ── 4. Proteção WAF ───────────────────────────────────────────────────
        _add_and_bookmark_heading(document, t("doc.headings.waf_protection_config", lang), 3, toc_list, counters)

        # Actions
        _add_and_bookmark_heading(document, t("doc.headings.waf_actions", lang), 4, toc_list, counters)
        if policy.actions:
            headers = [t("doc.common.name", lang), t("doc.common.type", lang), t("doc.common.code", lang)]
            table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
            _style_table_headers(table, headers)
            for a in policy.actions:
                cells = table.add_row().cells
                cells[0].text = a.name
                cells[1].text = a.type
                cells[2].text = str(a.code) if a.code else "N/A"
            document.add_paragraph()
        else:
            document.add_paragraph(t("doc.messages.no_config_found", lang))

        # Access Control
        _add_and_bookmark_heading(document, t("doc.headings.waf_access_control", lang), 4, toc_list, counters)
        if policy.access_control_rules:
            headers = [t("doc.common.name", lang), t("doc.headers.action", lang), t("doc.headers.condition", lang)]
            table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
            _style_table_headers(table, headers)
            for r in policy.access_control_rules:
                cells = table.add_row().cells
                cells[0].text = r.name
                cells[1].text = r.action_name
                cells[2].text = r.condition or "N/A"
            document.add_paragraph()
        else:
            document.add_paragraph(t("doc.messages.no_config_found", lang))

        # Rate Limiting
        _add_and_bookmark_heading(document, t("doc.headings.waf_rate_limiting", lang), 4, toc_list, counters)
        if policy.rate_limiting_rules:
            headers = [t("doc.common.name", lang), t("doc.headers.action", lang), t("doc.headers.condition", lang)]
            table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
            _style_table_headers(table, headers)
            for r in policy.rate_limiting_rules:
                cells = table.add_row().cells
                cells[0].text = r.name
                cells[1].text = r.action_name
                cells[2].text = r.condition or "N/A"
            document.add_paragraph()
        else:
            document.add_paragraph(t("doc.messages.no_config_found", lang))

        # Protection Rules
        _add_and_bookmark_heading(document, t("doc.headings.waf_protection_rules", lang), 4, toc_list, counters)
        if policy.protection_rules:
            headers = [t("doc.common.name", lang), t("doc.headers.action", lang), t("doc.headers.capabilities", lang)]
            table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
            _style_table_headers(table, headers)
            for r in policy.protection_rules:
                cells = table.add_row().cells
                cells[0].text = r.name
                cells[1].text = r.action_name
                caps = [f"{c.key} (v{c.version})" for c in r.protection_capabilities]
                cells[2].text = ", ".join(caps) if caps else "N/A"
            document.add_paragraph()
        else:
            document.add_paragraph(t("doc.messages.no_config_found", lang))


def _render_single_load_balancer(
    document: Document,
    lb: LoadBalancerData,
    infra_data: InfrastructureData,
    toc_list: list,
    counters: Dict[int, int],
    lang: str,
    base_level: int = 3,
):
    """Renders the standard details of a Load Balancer."""

    # ── General info ──────────────────────────────────────────────────────────
    ip_strings = [
        f"{ip.ip_address} ({t('doc.common.public', lang) if ip.is_public else t('doc.common.private', lang)})"
        for ip in lb.ip_addresses
    ]
    lb_info = {
        t("doc.common.name",   lang): lb.display_name,
        t("doc.headers.shape", lang): lb.shape_name,
        t("doc.common.state",  lang): lb.lifecycle_state,
        t("doc.headers.lb_ip_addresses", lang): "\n".join(ip_strings) or "N/A",
    }
    _create_titled_key_value_table(document, t("doc.headings.lb_general_info", lang), lb_info)

    # ── WAF Integration (if present) ──────────────────────────────────────────
    if getattr(lb, "waf_firewall_name", None) and getattr(lb, "waf_policy_name", None):
        waf_info = {
            t("doc.headings.waf_firewall",     lang): lb.waf_firewall_name,
            "Firewall OCID":                         lb.waf_firewall_id or "N/A",
            t("doc.headings.waf_policy_name",  lang): lb.waf_policy_name,
            "Policy OCID":                           lb.waf_policy_id  or "N/A",
        }
        _create_titled_key_value_table(document, t("doc.headings.waf_integration", lang), waf_info)

    # ── Hostnames virtuais ────────────────────────────────────────────────────
    if getattr(lb, "hostnames", None):
        _add_and_bookmark_heading(
            document, "Hostnames Virtuais", base_level + 1, toc_list, counters
        )
        if lb.hostnames:
            headers = [t("doc.common.name", lang)]
            table = document.add_table(rows=1, cols=1, style="Table Grid")
            _style_table_headers(table, headers)
            for h in lb.hostnames:
                table.add_row().cells[0].text = h.name
            document.add_paragraph()
        else:
            document.add_paragraph(t("doc.messages.no_resources_found", lang))

    # ── Listeners ─────────────────────────────────────────────────────────────
    _add_and_bookmark_heading(
        document, t("doc.headings.lb_listeners", lang), base_level + 1, toc_list, counters
    )
    if lb.listeners:
        # Build cert OCID → name map from infra_data
        certs_list   = list(getattr(infra_data, "certificates", []) or [])
        cert_id_map  = {c.get("id"): c.get("name") for c in certs_list if isinstance(c, dict) and c.get("id")}
        has_ssl      = any(
            (getattr(l, "ssl_certificate_ids", None) or [])
            for l in lb.listeners
        )

        headers = [
            t("doc.common.name",    lang),
            t("doc.common.protocol", lang),
            t("doc.common.port",    lang),
            t("doc.headers.lb_backend_set_default", lang),
        ]
        if has_ssl:
            headers.append("Certificado TLS")

        table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
        _style_table_headers(table, headers)

        for listener in lb.listeners:
            cells = table.add_row().cells
            cells[0].text = listener.name
            cells[1].text = listener.protocol
            cells[2].text = str(listener.port)
            cells[3].text = listener.default_backend_set_name
            if has_ssl:
                ssl_ids   = getattr(listener, "ssl_certificate_ids", None) or []
                cert_strs = [cert_id_map.get(oid, oid[-12:]) for oid in ssl_ids]
                cells[4].text = ", ".join(cert_strs) if cert_strs else "—"
    else:
        document.add_paragraph(t("doc.messages.no_listener_found", lang))
    document.add_paragraph()

    # ── Backend Sets ─────────────────────────────────────────────────────────
    _add_and_bookmark_heading(
        document, t("doc.headings.lb_backend_sets", lang), base_level + 1, toc_list, counters
    )
    if lb.backend_sets:
        for bs in lb.backend_sets:
            bs_info = {
                t("doc.headers.lb_backend_policy",  lang): bs.policy,
                t("doc.headers.lb_health_checker",  lang): f"{bs.health_checker.protocol}:{bs.health_checker.port} (Path: {bs.health_checker.url_path})",
            }
            _create_titled_key_value_table(
                document,
                f"{t('doc.headings.lb_backend_set_config', lang)}: {bs.name}",
                bs_info,
            )
            if bs.backends:
                headers = [
                    t("doc.headers.lb_backend_name",   lang),
                    t("doc.headers.lb_backend_ip",     lang),
                    t("doc.headers.lb_backend_port",   lang),
                    t("doc.headers.lb_backend_weight", lang),
                ]
                table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
                _style_table_headers(table, headers)
                for backend in bs.backends:
                    cells = table.add_row().cells
                    cells[0].text = backend.name
                    cells[1].text = backend.ip_address
                    cells[2].text = str(backend.port)
                    cells[3].text = str(backend.weight)
                document.add_paragraph()
            else:
                document.add_paragraph(t("doc.messages.no_backend_found", lang))
    else:
        document.add_paragraph(t("doc.messages.no_backend_set_found", lang))


def _add_compartment_certificates_section(
    document: Document,
    infra_data: InfrastructureData,
    toc_list: list,
    counters: Dict[int, int],
    lang: str,
) -> None:
    """
    Adds an OCI Certificates Service section to the infrastructure document.
    Lists Common Name, algorithms, validity window, and associations for each certificate.
    Only ACTIVE and PENDING_DELETION certificates are included.
    """
    certificates = [
        c for c in (getattr(infra_data, "certificates", []) or [])
        if isinstance(c, dict)
        and c.get("lifecycle_state", "").upper() in ("ACTIVE", "PENDING_DELETION")
    ]
    if not certificates:
        document.add_paragraph(t("doc.messages.no_certificates_found", lang))
        return

    # OCID → display_name index built from all LBs in scope.
    # Used to resolve the bound resource name in certificate associations,
    # since the association's display_name is the association object name,
    # not the LB name. lb.id is always present in LoadBalancerData.
    lb_ocid_to_name = {
        lb.id: lb.display_name
        for lb in (getattr(infra_data, "load_balancers", []) or [])
        if lb.id
    }

    for cert in sorted(certificates, key=lambda c: c.get("name", "")):
        name  = cert.get("name", "N/A")
        state = cert.get("lifecycle_state", "N/A")
        subj  = cert.get("subject", {}) or {}
        cvs   = cert.get("current_version_summary", {}) or {}
        assoc = cert.get("associations", []) or []

        _add_and_bookmark_heading(document, name, 2, toc_list, counters)

        cert_info = {
            t("doc.common.state",           lang): state,
            "Common Name":                         subj.get("common_name", "N/A"),
            t("doc.headers.cert_org",        lang): subj.get("organization", "N/A"),
            t("doc.headers.cert_key_algo",   lang): cert.get("key_algorithm", "N/A"),
            t("doc.headers.cert_sign_algo",  lang): cert.get("signature_algorithm", "N/A"),
            t("doc.headers.cert_stages",     lang): ", ".join(cvs.get("stages", [])) or "N/A",
            t("doc.headers.cert_valid_from", lang): cvs.get("valid_not_before", "N/A"),
            t("doc.headers.cert_valid_to",   lang): cvs.get("valid_not_after",  "N/A"),
        }
        if state == "PENDING_DELETION":
            cert_info[t("doc.headers.cert_deletion_date", lang)] = cert.get("time_of_deletion", "—")

        _create_titled_key_value_table(document, t("doc.headings.general_info", lang), cert_info)

        # Color the KV table rows based on certificate lifecycle state.
        _style_cert_kv_table_by_state(document, state)

        # SANs — separate table; apply row color directly on each row.
        sans = cert.get("subject_alternative_names", []) or []
        if sans:
            headers_san = ["SAN Type", "Value"]
            table_san = document.add_table(rows=1, cols=2, style="Table Grid")
            _style_table_headers(table_san, headers_san)
            for san in sans:
                if isinstance(san, dict):
                    cells = table_san.add_row().cells
                    cells[0].text = san.get("san_type", san.get("type", "N/A"))
                    cells[1].text = san.get("value", "N/A")
                    # Apply lifecycle state color to each SAN row.
                    _style_row_by_state(table_san.rows[-1], state)
            document.add_paragraph()

        # Certificate associations with other resources.
        # Name resolution order for the bound resource:
        #   1. Lookup by resource OCID in the LB index (most reliable)
        #   2. Fallback to display_name (association object name, less readable)
        #   3. Last resort: last 12 chars of the OCID
        if assoc:
            headers_a = [
                t("doc.common.name",  lang),
                t("doc.common.type",  lang),
                t("doc.common.state", lang),
                "OCID do Vínculo" if lang == "pt" else "Binding OCID",
            ]
            table_a = document.add_table(rows=1, cols=len(headers_a), style="Table Grid")
            _style_table_headers(table_a, headers_a)
            for a in assoc:
                if isinstance(a, dict):
                    resource_ocid = a.get("associated_resource_id", "")
                    resolved_name = (
                        lb_ocid_to_name.get(resource_ocid)
                        or a.get("display_name")
                        or (resource_ocid[-12:] if resource_ocid else "N/A")
                    )
                    assoc_state = a.get("lifecycle_state", "N/A")
                    cells = table_a.add_row().cells
                    cells[0].text = resolved_name
                    cells[1].text = a.get("resource_type", "N/A")
                    cells[2].text = assoc_state
                    cells[3].text = resource_ocid or "N/A"
                    _style_row_by_state(table_a.rows[-1], assoc_state)
            document.add_paragraph()